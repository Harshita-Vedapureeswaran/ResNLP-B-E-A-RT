# -*- coding: utf-8 -*-
"""capstone.ipynb
# FAKE DETECTION
"""

# Install necessary packages (if using Colab, run in a cell)
!pip install pdf2image pytesseract scikit-image pandas python-docx PyPDF2 sentence-transformers
!apt-get install -y poppler-utils tesseract-ocr

import os
import re
import cv2
import numpy as np
import pytesseract
from pytesseract import Output
from PIL import Image
import tensorflow as tf
from tensorflow.keras.layers import (Input, Conv2D, BatchNormalization, Activation,
                                     ZeroPadding2D, MaxPooling2D, AveragePooling2D,
                                     Flatten, Dense, Add, Lambda)
from tensorflow.keras.models import Model
from tensorflow.keras import backend as K
from pdf2image import convert_from_path
from skimage.metrics import structural_similarity as ssim
from matplotlib import pyplot as plt
from sentence_transformers import SentenceTransformer, util

import datetime
import pandas as pd
import docx
from datetime import datetime as dt
import PyPDF2
from sentence_transformers import SentenceTransformer, util
import difflib
import datetime
import PyPDF2
from docx import Document
import json
import difflib
import sys  # For potential early exit (not used now)

! pip install fitz
! pip install wordninja
!pip install PyPDF2 python-docx scikit-learn PyMuPDF

def crop_thomson_and_logo_expanded(image, target_tokens=None, offset_bottom=200,
                                   fallback_fraction=0.15, margin=100):
    """
    Crops a region that includes both the reference text (e.g., "©", "thomson", "reuters",
    "westlaw", "asia", "private", "limited") and the logo below it, with extra margin added.
    """
    if target_tokens is None:
        target_tokens = {"©", "thomson", "reuters", "westlaw", "asia", "private", "limited"}
    # Convert image to grayscale for OCR
    gray = cv2.cvtColor(image, cv2.COLOR_RGB2GRAY)
    data = pytesseract.image_to_data(gray, output_type=Output.DICT)
    #print("OCR detected texts:", data['text'])
    boxes = []
    n_boxes = len(data['text'])
    for i in range(n_boxes):
        token = data['text'][i].strip().lower()
        if token in target_tokens:
            x = data['left'][i]
            y = data['top'][i]
            w = data['width'][i]
            h = data['height'][i]
            boxes.append((x, y, x+w, y+h))
    H, W, _ = image.shape
    if boxes:
        min_x = max(min(box[0] for box in boxes) - margin, 0)
        min_y = max(min(box[1] for box in boxes) - margin, 0)
        max_x = min(max(box[2] for box in boxes) + margin, W)
        max_y = min(max(box[3] for box in boxes) + offset_bottom + margin, H)
        #print("Expanded bounding box: min_y =", min_y, "max_y =", max_y)
        cropped = image[min_y:max_y, min_x:max_x]
    else:
        crop_start_y = int(H * (1 - fallback_fraction))
        cropped = image[crop_start_y:, :]
    #print("Final cropped region shape:", cropped.shape)
    return cropped

#############################################
# Custom ResNet50-Like Network Components
#############################################
def identity_block(X, f, filters, stage, block):
    conv_name_base = f"res{stage}{block}_branch"
    bn_name_base = f"bn{stage}{block}_branch"
    F1, F2, F3 = filters

    X_shortcut = X
    X = Conv2D(F1, (1, 1), strides=(1, 1), padding='valid', name=conv_name_base + '2a')(X)
    X = BatchNormalization(axis=3, name=bn_name_base + '2a')(X)
    X = Activation('relu')(X)

    X = Conv2D(F2, (f, f), strides=(1, 1), padding='same', name=conv_name_base + '2b')(X)
    X = BatchNormalization(axis=3, name=bn_name_base + '2b')(X)
    X = Activation('relu')(X)

    X = Conv2D(F3, (1, 1), strides=(1, 1), padding='valid', name=conv_name_base + '2c')(X)
    X = BatchNormalization(axis=3, name=bn_name_base + '2c')(X)

    X = Add()([X, X_shortcut])
    X = Activation('relu')(X)
    return X

def convolutional_block(X, f, filters, stage, block, s=2):
    conv_name_base = f"res{stage}{block}_branch"
    bn_name_base = f"bn{stage}{block}_branch"
    F1, F2, F3 = filters

    X_shortcut = X
    X = Conv2D(F1, (1, 1), strides=(s, s), padding='valid', name=conv_name_base + '2a')(X)
    X = BatchNormalization(axis=3, name=bn_name_base + '2a')(X)
    X = Activation('relu')(X)

    X = Conv2D(F2, (f, f), strides=(1, 1), padding='same', name=conv_name_base + '2b')(X)
    X = BatchNormalization(axis=3, name=bn_name_base + '2b')(X)
    X = Activation('relu')(X)

    X = Conv2D(F3, (1, 1), strides=(1, 1), padding='valid', name=conv_name_base + '2c')(X)
    X = BatchNormalization(axis=3, name=bn_name_base + '2c')(X)

    X_shortcut = Conv2D(F3, (1, 1), strides=(s, s), padding='valid', name=conv_name_base + '1')(X_shortcut)
    X_shortcut = BatchNormalization(axis=3, name=bn_name_base + '1')(X_shortcut)

    X = Add()([X, X_shortcut])
    X = Activation('relu')(X)
    return X

def create_custom_resnet50(input_shape=(224, 224, 3)):
    X_input = Input(input_shape)

    # Stage 1
    X = ZeroPadding2D((3, 3))(X_input)
    X = Conv2D(64, (7, 7), strides=(2, 2), name='conv1')(X)
    X = BatchNormalization(axis=3, name='bn_conv1')(X)
    X = Activation('relu')(X)
    X = ZeroPadding2D((1, 1))(X)
    X = MaxPooling2D((3, 3), strides=(2, 2))(X)

    # Stage 2
    X = convolutional_block(X, f=3, filters=[64, 64, 256], stage=2, block='a', s=1)
    X = identity_block(X, 3, [64, 64, 256], stage=2, block='b')
    X = identity_block(X, 3, [64, 64, 256], stage=2, block='c')

    # Stage 3
    X = convolutional_block(X, f=3, filters=[128, 128, 512], stage=3, block='a', s=2)
    X = identity_block(X, 3, [128, 128, 512], stage=3, block='b')
    X = identity_block(X, 3, [128, 128, 512], stage=3, block='c')
    X = identity_block(X, 3, [128, 128, 512], stage=3, block='d')

    # Stage 4
    X = convolutional_block(X, f=3, filters=[256, 256, 1024], stage=4, block='a', s=2)
    X = identity_block(X, 3, [256, 256, 1024], stage=4, block='b')
    X = identity_block(X, 3, [256, 256, 1024], stage=4, block='c')
    X = identity_block(X, 3, [256, 256, 1024], stage=4, block='d')
    X = identity_block(X, 3, [256, 256, 1024], stage=4, block='e')
    X = identity_block(X, 3, [256, 256, 1024], stage=4, block='f')

    # Stage 5
    X = convolutional_block(X, f=3, filters=[512, 512, 2048], stage=5, block='a', s=2)
    X = identity_block(X, 3, [512, 512, 2048], stage=5, block='b')
    X = identity_block(X, 3, [512, 512, 2048], stage=5, block='c')

    # Average Pooling and Feature Vector
    X = AveragePooling2D(pool_size=(7, 7), name='avg_pool')(X)
    X = Flatten()(X)
    X = Dense(256, activation='relu')(X)

    model = Model(inputs=X_input, outputs=X, name='Custom_ResNet50')
    return model

#############################################
# Utility Functions for Logo Verification
#############################################
def center_crop(image, crop_width_fraction=0.9):
    h, w, _ = image.shape
    target_w = int(w * crop_width_fraction)
    start_x = (w - target_w) // 2
    return image[:, start_x:start_x+target_w]

def mse(imageA, imageB):
    err = np.sum((imageA.astype("float") - imageB.astype("float")) ** 2)
    err /= float(imageA.shape[0] * imageA.shape[1])
    return err

#############################################
# Build Siamese Network for Logo Comparison
#############################################
def build_siamese_network():
    input_logo_test = Input(shape=(224, 224, 3), name="logo_test")
    input_logo_original = Input(shape=(224, 224, 3), name="logo_original")
    base_network = create_custom_resnet50((224, 224, 3))
    feat_test = base_network(input_logo_test)
    feat_original = base_network(input_logo_original)
    def euclidean_distance(vects):
        x, y = vects
        sum_square = K.sum(K.square(x - y), axis=1, keepdims=True)
        return K.sqrt(K.maximum(sum_square, K.epsilon()))
    distance = Lambda(euclidean_distance, name="euclidean_distance")([feat_test, feat_original])
    output = Dense(1, activation='sigmoid', name="authenticity_output")(distance)
    model = Model(inputs=[input_logo_test, input_logo_original], outputs=output)
    model.compile(optimizer='adam', loss='binary_crossentropy', metrics=['accuracy'])
    return model

#############################################
# Main Verification Function
#############################################
def verify_document_structure(test_path,original_path):
    """
    Runs all steps (logo-based verification and SSIM/MSE similarity) in one function.
    Updates the global variable 'document_fake' accordingly.
    Returns True if the document appears authentic, or False if fake.
    """
    document_fake=False
    # ----- Logo Verification -----
    # Convert PDFs to images (assuming test and original logo PDFs are provided)
    pages_test = convert_from_path(test_path, dpi=300)
    pages_original = convert_from_path(original_path, dpi=300)

    # Assume logos are on the last page.
    doc_test = np.array(pages_test[-1])
    doc_original = np.array(pages_original[-1])

    # Crop logos using OCR-based cropping function.
    cropped_region_test = crop_thomson_and_logo_expanded(doc_test,
                                                         target_tokens={"©", "thomson", "reuters", "westlaw", "asia", "private", "limited"},
                                                         offset_bottom=200,
                                                         fallback_fraction=0.15,
                                                         margin=100)
    cropped_region_original = crop_thomson_and_logo_expanded(doc_original,
                                                             target_tokens={"©", "thomson", "reuters", "westlaw", "asia", "private", "limited"},
                                                             offset_bottom=200,
                                                             fallback_fraction=0.15,
                                                             margin=100)
    # Optionally center-crop horizontally.
    cropped_region_test = center_crop(cropped_region_test, crop_width_fraction=0.9)
    cropped_region_original = center_crop(cropped_region_original, crop_width_fraction=0.9)

    # Resize cropped regions to 224x224.
    logo_test_resized = cv2.resize(cropped_region_test, (224, 224))
    logo_original_resized = cv2.resize(cropped_region_original, (224, 224))

    # Save images for inspection.
    Image.fromarray(cv2.cvtColor(logo_test_resized, cv2.COLOR_BGR2RGB)).save("cropped_logo_test.png")
    Image.fromarray(cv2.cvtColor(logo_original_resized, cv2.COLOR_BGR2RGB)).save("cropped_logo_original.png")
    #print("Cropped logos saved as 'cropped_logo_test.png' and 'cropped_logo_original.png'.")

    # Normalize and add batch dimension.
    logo_test_norm = logo_test_resized.astype(np.float32) / 255.0
    logo_original_norm = logo_original_resized.astype(np.float32) / 255.0
    logo_test_batch = np.expand_dims(logo_test_norm, axis=0)
    logo_original_batch = np.expand_dims(logo_original_norm, axis=0)
    print("\n========== Document Structure Analysis ==========")
    # Build Siamese network and get prediction.
    siamese_model = build_siamese_network()
    siamese_prediction = siamese_model.predict([logo_test_batch, logo_original_batch])
    print("Siamese network prediction (ResNet50-based):", siamese_prediction[0][0])

    # Compute SSIM and MSE metrics.
    gray_test = cv2.cvtColor(logo_test_resized, cv2.COLOR_BGR2GRAY)
    gray_original = cv2.cvtColor(logo_original_resized, cv2.COLOR_BGR2GRAY)
    ssim_score, diff = ssim(gray_test, gray_original, full=True)
    print("Structural Similarity Index Metric score:", ssim_score)
    def mse(imageA, imageB):
        err = np.sum((imageA.astype("float") - imageB.astype("float")) ** 2)
        err /= float(imageA.shape[0] * imageA.shape[1])
        return err
    mse_value = mse(gray_test, gray_original)
    print("MSE:", mse_value)

    # Combine predictions (for example, average of SSIM and Siamese prediction).
    combined_probability = (0.7*ssim_score + 0.3*siamese_prediction[0][0])
    print("Combined authenticity probability:", combined_probability)

    # Update document_fake flag based on threshold.
    if combined_probability < 0.70:
        print("Document is Fake due to low authenticity probability.")
        document_fake = True
    else:
        print("Document appears Authentic based on logo similarity metrics.")
    return not document_fake

import os
import re
import datetime
import pandas as pd
import docx
from datetime import datetime as dt
import PyPDF2
from sentence_transformers import SentenceTransformer, util
import difflib

#############################################
# PART 1: INVENTORY LOADING
#############################################
def load_inventory_from_excel(excel_file):
    """
    Loads the inventory of known authentic cases from an Excel file
    with columns:
      Case Name, Court, Date, Year, Subject, Where Reported, Documents.
    Returns a list of dictionaries (one per row).
    """
    df = pd.read_excel(excel_file)
    return df.to_dict(orient="records")

#############################################
# PART 2: DOCUMENT EXTRACTION
#############################################
def extract_text_from_docx(docx_file):
    """
    Reads a .docx file and returns a list of non-empty paragraph strings.
    """
    doc = docx.Document(docx_file)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return paragraphs

def extract_text_from_file(file_path):
    """
    Extracts text from a PDF file using PyPDF2.
    Returns the extracted text as a single string.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + " "
        return text
    elif ext in [".docx", ".doc"]:
        return "\n".join(extract_text_from_docx(file_path))
    else:
        raise ValueError("Unsupported file type: " + ext)

def preprocess_text(text):
    """
    Lowercases the text and condenses whitespace.
    """
    text = text.lower()
    text = re.sub(r'\s+', ' ', text).strip()
    return text
def parse_delhi_document_fields(paragraphs):
    """
    Parses a document (e.g., from RTDoc 16.docx) with lines like:
      Number of documents delivered : 1
      Delhi High Court
      10 November 2027
      Sara Carrierre Dubey
      v
      Ashish Dubey
      Case No : CRL. M. C. 574/2020 ...
      Bench : ...
      Citation : 2025 Indlaw DEL ...
    Extracts Court, Date, Case Name, Case No, and Citation.
    """
    record = {
        "Court": "",
        "Date": "",
        "Case Name": "",
        "Case No": "",
        "Citation": ""
    }

    # Find index after "Number of documents delivered"
    start_index = 0
    for i, line in enumerate(paragraphs):
        if "Number of documents delivered" in line:
            start_index = i + 1
            break

    # Next line is Court, then next line is Date
    if start_index < len(paragraphs):
        record["Court"] = paragraphs[start_index].strip()
    if start_index + 1 < len(paragraphs):
        record["Date"] = paragraphs[start_index + 1].strip()

    # Collect case name lines until we hit "Case No :"
    case_name_lines = []
    i = start_index + 2
    while i < len(paragraphs):
        line = paragraphs[i].strip()
        if line.startswith("Case No :"):
            break
        case_name_lines.append(line)
        i += 1
    record["Case Name"] = " ".join(case_name_lines).strip()

    # Find "Case No :"
    for line in paragraphs:
        if line.startswith("Case No :"):
            record["Case No"] = line.replace("Case No :", "").strip()
            break

    # Find "Citation :"
    for line in paragraphs:
        if line.startswith("Citation :"):
            record["Citation"] = line.replace("Citation :", "").strip()
            break

    return record

#############################################
# PART 3: DOCUMENT FIELD PARSING VIA REGEX
#############################################
def parse_delhi_document_fields_from_text(text):
    """
    Uses a regular expression to extract document fields from the entire text.
    Expected format (with all words joined):

    "Supreme Court of India 23 April 2020 Kamal Parti v Raj Kumar Parti and another
     Case No : CS (OS) 191/2016 & IAs No. 5040/2016 (u/O XXXIX R-1&2 CPC), 12979/2016 (u/S 151 CPC) & 3048/2018 (u/O VIII R-1A CPC)
     Bench : Rajiv Sahai Endlaw
     Citation : 2027 Indlaw KAR 915"

    The regex stops the Citation field when it sees (case-insensitively) "order", "the order", "summary",
    or "the summary" as a separate word, or the end.

    Returns a dictionary with keys: Court, Date, Case Name, Case No, Bench, Citation.
    """
    pattern = re.compile(
        r'(?P<court>.+?)\s+'
        r'(?P<date>\d{1,2}\s+[A-Za-z]+\s+\d{4})\s+'
        r'(?P<plaintiff>.+?)\s+v\s+'
        r'(?P<defendant>.+?)\s+'
        r'Case\s*No\s*:\s*(?P<caseno>.+?)\s+'
        r'Bench\s*:\s*(?P<bench>.+?)\s+'
        r'Citation\s*:\s*(?P<citation>.+?)(?=\s+(?:the\s+)?(?:order|summary|judgement|judgment)\b|$)',
        re.IGNORECASE | re.DOTALL
    )
    m = pattern.search(text)
    if m:
        record = {}
        record["Court"] = re.sub(r'\s+', ' ', m.group("court")).strip()
        record["Date"] = re.sub(r'\s+', ' ', m.group("date")).strip()
        plaintiff = re.sub(r'\s+', ' ', m.group("plaintiff")).strip()
        defendant = re.sub(r'\s+', ' ', m.group("defendant")).strip()
        record["Case Name"] = f"{plaintiff} v {defendant}"
        record["Case No"] = re.sub(r'\s+', ' ', m.group("caseno")).strip()
        record["Bench"] = re.sub(r'\s+', ' ', m.group("bench")).strip()
        record["Citation"] = re.sub(r'\s+', ' ', m.group("citation")).strip()

        return record
    else:
        print("Regex did not match the expected pattern.")
        return {}

#############################################
# PART 4: MAPPING & FLEXIBLE MATCH
#############################################
def normalize_date(date_str):
    """
    Tries to parse various date formats (e.g., "23 April 2020" or "10-Nov-20")
    into a standard YYYY-MM-DD string.
    Returns None if parsing fails.
    """
    date_formats = [
        "%d %B %Y",
        "%d-%b-%y",
        "%d-%m-%Y",
        "%d/%m/%Y",
        "%d %b %Y"
    ]
    date_str_clean = date_str.strip().rstrip(".,")
    for fmt in date_formats:
        try:
            dt_obj = dt.strptime(date_str_clean, fmt)
            return dt_obj.strftime("%Y-%m-%d")
        except ValueError:
            continue
    return None
def process_court_field(court_field):
    """
    Processes the court field:
      - If the field contains the phrase "court of india" (case-insensitive), returns the entire field.
      - Otherwise, returns the text up to and including the first occurrence of the word "court".
    """
    if not court_field:
        return ""
    if re.search(r'court of india', court_field, re.IGNORECASE):
        return court_field.strip()
    else:
        match = re.search(r'\bcourt\b', court_field, re.IGNORECASE)
        if match:
            return court_field[:match.end()].strip()
        return court_field.strip()
    return court_field.strip()
def clean_text(s):
    """
    Lowercases, condenses all whitespace (including newlines) to a single space,
    strips leading/trailing whitespace, and removes trailing punctuation.
    Converts non-string values to a string.
    """
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)
    s = s.lower()
    s = re.sub(r'\s+', ' ', s)  # Replace all whitespace sequences with a single space
    s = s.strip()
    s = re.sub(r'\s+([,;:])', r'\1', s)  # Remove extra spaces before punctuation
    return s.rstrip(".,")

def map_extracted_to_inventory_format(extracted):
    """
    Maps fields from the document to the inventory format:
      Case Name, Court, Date, Year, Subject, Where Reported, Documents.
    """
    mapped = {}
    mapped["Case Name"] = extracted.get("Case Name", "")
    mapped["Court"] = process_court_field(extracted.get("Court", ""))
    original_date_str = extracted.get("Date", "")
    iso_date = normalize_date(original_date_str)
    mapped["Date"] = iso_date if iso_date else original_date_str
    if iso_date:
        mapped["Year"] = iso_date.split("-")[0]
    else:
        mapped["Year"] = ""
    mapped["Subject"] = ""
    mapped["Where Reported"] = extracted.get("Citation", "")
    mapped["Documents"] = ""
    return mapped

def flexible_check(mapped_extracted, inventory):
    """
    Performs a flexible check:
      1) Filters inventory records by the same Year (if available).
      2) Further filters by the citation field ("Where Reported").
      3) Compares Case Name, Court, Date, and Where Reported in a case-insensitive manner.
    Returns True if a matching record is found, else False.
    If a match is found, prints the matching inventory record.
    """
    doc_year = mapped_extracted.get("Year", "").strip()
    if doc_year:
        inv_filtered = [item for item in inventory if str(item.get("Year", "")).strip() == doc_year]
    else:
        inv_filtered = inventory

    # Further filter by citation.
    doc_citation = clean_text(mapped_extracted.get("Where Reported", ""))
    if doc_citation:
        inv_filtered = [item for item in inv_filtered if clean_text(item.get("Where Reported", "")) == doc_citation]

    doc_case_name = clean_text(mapped_extracted["Case Name"])
    doc_court = clean_text(mapped_extracted["Court"])
    doc_date_iso = mapped_extracted["Date"]
    doc_where = doc_citation
    #print(inv_filtered)
    for item in inv_filtered:
        inv_case_name = clean_text(item.get("Case Name", ""))
        inv_court = clean_text(item.get("Court", ""))
        inv_date_val = item.get("Date", "")
        if isinstance(inv_date_val, (datetime.datetime, datetime.date)):
            inv_date_str = inv_date_val.strftime("%Y-%m-%d")
        else:
            inv_date_str = str(inv_date_val).strip()
        inv_date_iso = normalize_date(inv_date_str) or inv_date_str
        inv_where = clean_text(item.get("Where Reported", ""))

        if doc_case_name != inv_case_name:
            continue
        if doc_court != inv_court:
            continue
        try:
            doc_date_obj = dt.strptime(doc_date_iso, "%Y-%m-%d")
            inv_date_obj = dt.strptime(inv_date_iso, "%Y-%m-%d")
            if doc_date_obj != inv_date_obj:
                continue
        except Exception:
            if clean_text(doc_date_iso) != clean_text(inv_date_iso):
                continue
        if doc_where != inv_where:
            continue
        #print("\nMatched inventory record:")
        #print(item)
        return True

    return False

#############################################
# PART 5: MAIN VERIFICATION FUNCTION
#############################################
def verify_case_information(file_path):
    """
    Verifies a document's case information against the inventory.
    Works for both PDF and DOCX files.
    Returns True if a matching record is found (document is authentic), else False.
    """
    inventory_file = "cases-- (2).xlsx"  # Update this path as needed
    inventory = load_inventory_from_excel(inventory_file)

    ext = os.path.splitext(file_path)[1].lower()
    if ext in [".docx", ".doc"]:
        paragraphs = extract_text_from_docx(file_path)
        full_text = " ".join(paragraphs)
    elif ext == ".pdf":
        full_text = extract_text_from_file(file_path)
    else:
        raise ValueError("Unsupported file type: " + ext)

    full_text = full_text.strip()

    # Use regex-based parsing on the full text.
    extracted = parse_delhi_document_fields_from_text(full_text)
    if not extracted:
        # Fallback: try splitting into lines and using the original parser.
        lines = [line.strip() for line in full_text.splitlines() if line.strip() != ""]
        extracted = parse_delhi_document_fields(lines)
    print("\n========== Case Details Checker ==========")
    print("Extracted fields from  document:")
    for k, v in extracted.items():
        print(f"{k}: {v}")
    document_fake=False
    mapped = map_extracted_to_inventory_format(extracted)
    #print("\nMapped fields for authenticity check:")

    is_authentic = flexible_check(mapped, inventory)
    if is_authentic:
        print("\nThis document MATCHES a record in the inventory. It appears AUTHENTIC.")
    else:
        print("\nNo exact match found in the inventory. Document may NOT be authentic.")
        document_fake = True

    return not document_fake

# End of code.

def extract_text_from_file(file_path, num_pages=None):
    """
    Extracts text from a PDF or DOCX file.
    If num_pages is specified for a PDF, only that many pages are read.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            pages_to_read = len(pdf_reader.pages) if num_pages is None else min(num_pages, len(pdf_reader.pages))
            for i in range(pages_to_read):
                page_text = pdf_reader.pages[i].extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    elif ext == ".docx":
        document = Document(file_path)
        return "\n".join([para.text for para in document.paragraphs])
    else:
        raise ValueError("Unsupported file type: " + ext)

def preprocess_text(text):
    """
    Lowercases the text and condenses whitespace.
    """
    text = text.lower()
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def extract_header_date(text, max_lines=10):
    """
    Attempts to extract a header date from the first few non-empty lines.
    First, it checks each line individually; if no complete date is found,
    it will try concatenating lines 3-7 and search again.
    Returns a datetime.date object if found, else None.
    """
    lines = [line.strip() for line in text.splitlines() if line.strip() != ""]
    # Try each of the first max_lines lines individually.
    for i in range(min(max_lines, len(lines))):
        match = re.search(r'\b(\d{1,2}\s+[A-Za-z]+\s+\d{4})\b', lines[i])
        if match:
            try:
                return datetime.datetime.strptime(match.group(1), "%d %B %Y").date()
            except ValueError:
                continue
    # If not found, try concatenating lines 3 to 7.
    if len(lines) >= 7:
        combined = " ".join(lines[3:7])
        match = re.search(r'\b(\d{1,2}\s+[A-Za-z]+\s+\d{4})\b', combined)
        if match:
            try:
                return datetime.datetime.strptime(match.group(1), "%d %B %Y").date()
            except ValueError:
                pass
    return None

def extract_citation_year(text):
    """
    Searches the entire text for a pattern like 'Citation: 2015' or 'Citation - 2015'
    (case-insensitive) and extracts the first four-digit year encountered.
    Returns the year as an integer, or None if no such pattern is found.
    """
    pattern = re.compile(r'citation\s*[:\-]\s*(\d{4})', re.IGNORECASE)
    match = pattern.search(text)
    if match:
        try:
            return int(match.group(1))
        except ValueError:
            return None
    return None

def verify_document_date(file_path):
    """
    Verifies the document's date information.
      1. Extracts full text from the document.
      2. Attempts to extract a header date from the first few lines.
      3. Checks if the header date is in the future relative to today.
      4. Extracts a citation year from the document.
      5. Compares the header date's year with the citation year.
    Prints detailed results and returns True if the document date is valid,
    or False if discrepancies are found.
    """
      # Reset global flag
    document_fake=False
    # Extract text and preprocess
    raw_text = extract_text_from_file(file_path)
    preprocessed_text = preprocess_text(raw_text)

    # Extract header date from the raw text (not preprocessed, to preserve case)
    header_date = extract_header_date(raw_text, max_lines=10)
    today = datetime.date.today()

    if header_date:
        print(f"Extracted header date: {header_date.strftime('%d %B %Y')}")
        if header_date > today:
            print(f"Header date {header_date} is in the future relative to today ({today}).")
            document_fake = True
        else:
            print(f"")
    else:
        print("No header date found in the document.")
        document_fake = True

    # Extract citation year from the raw text.
    citation_year = extract_citation_year(raw_text)
    if citation_year:
        print(f"Citation year extracted: {citation_year}")
    else:
        #print("No citation year found in the document.")
        document_fake = True

    # Compare header date's year with citation year if both exist.
    if header_date and citation_year:
        header_year = header_date.year
        if header_year == citation_year:
            print(f"Header year ({header_year}) matches the citation year ({citation_year}).")
        else:
            print(f"Header year ({header_year}) does NOT match the citation year ({citation_year}).")
            document_fake = True

    print("\n========== Document Date Analysis ==========")
    if document_fake:
        print("🚨 Document is FAKE due to date discrepancies!")

    else:
        print("✅ Document date appears valid.")
    return not document_fake

def extract_text_from_pdf(file_path, num_pages):
    """
    Extracts text from the first `num_pages` pages of a PDF.
    """
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        total_pages = len(pdf_reader.pages)
        pages_to_read = min(num_pages, total_pages)
        for i in range(pages_to_read):
            page_text = pdf_reader.pages[i].extract_text()
            if page_text:
                text += page_text + "\n"
    return text

# Specify the PDF file containing the sections.
pdf_file = 'IPC.pdf'
pdf_text = extract_text_from_pdf(pdf_file, num_pages=13)

# Define a regex pattern to capture lines like "1. Title and extent ...."
# The pattern captures a leading number followed by a period and whitespace, then the rest of the line.
pattern = re.compile(r'^(\d+)\.\s+(.*)$')

# Split the extracted text into lines and filter only those that match our pattern.
lines = pdf_text.splitlines()
sections_dict = {}

for line in lines:
    line = line.strip()
    # Skip header lines such as "SECTIONS"
    if line.upper() == "SECTIONS":
        continue
    match = pattern.match(line)
    if match:
        section_number = match.group(1)
        section_title = match.group(2)
        sections_dict[section_number] = section_title

# Save the dictionary to a JSON file using readable Unicode.
with open('sections_dict.json', 'w', encoding='utf-8') as f:
    json.dump(sections_dict, f, indent=4, ensure_ascii=False)

print("Extraction complete. Dictionary saved to sections_dict.json")

# Specify the PDF file containing the sections.
pdf_file = 'CrPC.pdf'
pdf_text = extract_text_from_pdf(pdf_file, num_pages=20)

# Define a regex pattern to capture lines like "1. Title and extent ...."
# The pattern captures a leading number followed by a period and whitespace, then the rest of the line.
pattern = re.compile(r'^(\d+)\.\s+(.*)$')

# Split the extracted text into lines and filter only those that match our pattern.
lines = pdf_text.splitlines()
sections_dict = {}

for line in lines:
    line = line.strip()
    # Skip header lines such as "SECTIONS"
    if line.upper() == "SECTIONS":
        continue
    match = pattern.match(line)
    if match:
        section_number = match.group(1)
        section_title = match.group(2)
        sections_dict[section_number] = section_title

# Save the dictionary to a JSON file using readable Unicode.
with open('crpc_dict.json', 'w', encoding='utf-8') as f:
    json.dump(sections_dict, f, indent=4, ensure_ascii=False)

print("Extraction complete. Dictionary saved to sections_dict.json")

# Global flag for document authenticity
document_fake = False

# ---------- Text Extraction Functions ----------
def extract_text_from_file(file_path, num_pages=None):
    """
    Extracts text from a PDF or DOCX file.
    If num_pages is specified for a PDF, only that many pages are read.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            pages_to_read = len(pdf_reader.pages) if num_pages is None else min(num_pages, len(pdf_reader.pages))
            for i in range(pages_to_read):
                page_text = pdf_reader.pages[i].extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    elif ext == ".docx":
        document = Document(file_path)
        return "\n".join([para.text for para in document.paragraphs])
    else:
        raise ValueError("Unsupported file type: " + ext)

def preprocess_text(text):
    """
    Normalizes text by lowercasing and condensing whitespace.
    """
    text = text.lower()
    text = re.sub(r'\s+', ' ', text).strip()
    return text

# ---------- Load Legal Reference Dictionaries ----------
def load_legal_resources(ipc_path='sections_dict.json', crpc_path='crpc_dict.json', model_name='nlpaueb/legal-bert-base-uncased'):
    """
    Loads legal reference dictionaries and the SentenceTransformer model.
    """
    with open(ipc_path, 'r', encoding='utf-8') as f:
        ipc_dict = json.load(f)
    with open(crpc_path, 'r', encoding='utf-8') as f:
        crpc_dict = json.load(f)
    valid_ipc_sections = set(ipc_dict.keys())
    valid_crpc_sections = set(crpc_dict.keys())
    try:
        st_model = SentenceTransformer(model_name)
        #print("Loaded legal-specific model.")
    except Exception as e:
        #print("Legal-specific model not available, using general-purpose model. Error:", e)
        st_model = SentenceTransformer('all-MiniLM-L6-v2')
    SIMILARITY_THRESHOLD = 0.58
    OVERALL_THRESHOLD = 0.58
    return ipc_dict, crpc_dict, valid_ipc_sections, valid_crpc_sections, st_model, SIMILARITY_THRESHOLD, OVERALL_THRESHOLD

# ---------- Similarity Functions ----------
def fuzzy_similarity(str1, str2):
    """Compute fuzzy similarity using difflib's SequenceMatcher."""
    return difflib.SequenceMatcher(None, str1, str2).ratio()

def check_context_usage(sentence, section, label, ipc_dict, crpc_dict, st_model):
    """
    For a given sentence and section number (as string), look up the reference definition
    from the appropriate dictionary (IPC or Cr.P.C) and compute:
      - Cosine similarity (using SentenceTransformer embeddings)
      - Fuzzy similarity (using difflib)
    Returns a tuple: (cosine_score, fuzzy_score)
    """
    if label == "I.P.C":
        reference_text = ipc_dict.get(section, "")
    else:
        reference_text = crpc_dict.get(section, "")
    if not reference_text:
        return None, None
    sentence_emb = st_model.encode(sentence, convert_to_tensor=True)
    ref_emb = st_model.encode(reference_text, convert_to_tensor=True)
    cosine_score = util.pytorch_cos_sim(sentence_emb, ref_emb).item()
    fuzzy_score = fuzzy_similarity(sentence, reference_text)
    return cosine_score, fuzzy_score

# ---------- Main Citation Verification Loop ----------
def section_checker(file_path):
    """
    Checks citation usage in the document.
    Returns True if document appears authentic; else False.
    If no sections are extracted, document is assumed authentic.
    """

    document_fake = False
    ipc_path='sections_dict.json'
    crpc_path='crpc_dict.json'
    model_name='nlpaueb/legal-bert-base-uncased'
    with open(ipc_path, 'r', encoding='utf-8') as f:
        ipc_dict = json.load(f)
    with open(crpc_path, 'r', encoding='utf-8') as f:
        crpc_dict = json.load(f)
    valid_ipc_sections = set(ipc_dict.keys())
    valid_crpc_sections = set(crpc_dict.keys())

    try:
        st_model = SentenceTransformer(model_name)
        #print("Loaded legal-specific model.")
    except Exception as e:
        #print("Legal-specific model not available, using general-purpose model. Error:", e)
        st_model = SentenceTransformer('all-MiniLM-L6-v2')
    SIMILARITY_THRESHOLD = 0.58
    OVERALL_THRESHOLD = 0.58


    # Extract and preprocess text.
    uploaded_text = extract_text_from_file(file_path)
    uploaded_text = preprocess_text(uploaded_text)

    def skip_generic_act_sentence(sentence):
        """
        Skip sentences with a generic "section ... of ... act" pattern unless they explicitly mention IPC or Cr.P.C.
        """
        if re.search(r'section\s+\S+\s+of\s+\S+', sentence, re.IGNORECASE):
            if not (re.search(r'\b(?:ipc|indian penal code|criminal penal code)\b', sentence, re.IGNORECASE) or
                    re.search(r'\b(?:cr\.?p\.?c|the code of criminal procedure|criminal procedure code)\b', sentence, re.IGNORECASE)):
                return True
        return False

    sentences = re.split(r'(?<=[.!?])\s+', uploaded_text)
    all_combined_scores = []

    for sentence in sentences:
        if skip_generic_act_sentence(sentence):
            continue
        if re.search(r'\b(?:d\.?v\.?(?:\s+act)?|domestic violence act|protection act)\b', sentence, re.IGNORECASE):
            continue
        if not re.search(r'\bSections?\b', sentence, re.IGNORECASE):
            continue

        explicit_ipc = re.search(r'\b(?:ipc|indian penal code|criminal penal code)\b', sentence, re.IGNORECASE)
        explicit_crpc = re.search(r'\b(?:cr\.?p\.?c|the code of criminal procedure|criminal procedure code)\b', sentence, re.IGNORECASE)
        if not (explicit_ipc or explicit_crpc):
            continue

        extracted_sections = set()
        singular = re.findall(r'Section\s+(\d+)', sentence, re.IGNORECASE)
        extracted_sections.update(singular)
        plural = re.findall(r'Sections?\s+([0-9,\sandalso]+)', sentence, re.IGNORECASE)
        for match in plural:
            numbers = re.findall(r'\d+', match)
            extracted_sections.update(numbers)
        range_matches = re.findall(r'Sections?\s+(\d+)\s*(?:-|to)\s*(\d+)', sentence, re.IGNORECASE)
        for start, end in range_matches:
            for num in range(int(start), int(end) + 1):
                extracted_sections.add(str(num))

        if explicit_ipc and not explicit_crpc:
            chosen_system = "I.P.C"
            labels_to_check = ["I.P.C"]
        elif explicit_crpc and not explicit_ipc:
            chosen_system = "Cr.P.C"
            labels_to_check = ["Cr.P.C"]
        else:
            system_votes = {"I.P.C": [], "Cr.P.C": []}
            for sec in extracted_sections:
                if sec in valid_ipc_sections:
                    cosine_ipc, fuzzy_ipc = check_context_usage(sentence, sec, "I.P.C", ipc_dict, crpc_dict, st_model)
                    if cosine_ipc is not None and fuzzy_ipc is not None:
                        ipc_score = 0.7 * cosine_ipc + 0.3 * fuzzy_ipc
                        system_votes["I.P.C"].append(ipc_score)
                if sec in valid_crpc_sections:
                    cosine_crpc, fuzzy_crpc = check_context_usage(sentence, sec, "Cr.P.C", ipc_dict, crpc_dict, st_model)
                    if cosine_crpc is not None and fuzzy_crpc is not None:
                        crpc_score = 0.7 * cosine_crpc + 0.3 * fuzzy_crpc
                        system_votes["Cr.P.C"].append(crpc_score)
            avg_ipc = sum(system_votes["I.P.C"]) / len(system_votes["I.P.C"]) if system_votes["I.P.C"] else 0
            avg_crpc = sum(system_votes["Cr.P.C"]) / len(system_votes["Cr.P.C"]) if system_votes["Cr.P.C"] else 0
            if avg_ipc >= avg_crpc:
                chosen_system = "I.P.C"
                labels_to_check = ["I.P.C"]
            else:
                chosen_system = "Cr.P.C"
                labels_to_check = ["Cr.P.C"]

        valid_candidates = []
        for sec in extracted_sections:
            candidate_found = False
            best_combined = -1.0
            best_detail = ""
            for label in labels_to_check:
                if label == "I.P.C":
                    if sec in valid_ipc_sections:
                        candidate_found = True
                        detail = ipc_dict.get(sec, 'Not found')
                        if detail != 'Not found':
                            cosine_score, fuzzy_score = check_context_usage(sentence, sec, "I.P.C", ipc_dict, crpc_dict, st_model)
                            if cosine_score is None or fuzzy_score is None:
                                continue
                            combined_score = 0.7 * cosine_score + 0.3 * fuzzy_score
                            if combined_score > best_combined:
                                best_combined = combined_score
                                best_detail = f"(I.P.C): {detail} (Combined Score: {combined_score:.2f})"
                elif label == "Cr.P.C":
                    if sec in valid_crpc_sections:
                        candidate_found = True
                        detail = crpc_dict.get(sec, 'Not found')
                        if detail != 'Not found':
                            cosine_score, fuzzy_score = check_context_usage(sentence, sec, "Cr.P.C", ipc_dict, crpc_dict, st_model)
                            if cosine_score is None or fuzzy_score is None:
                                continue
                            combined_score = 0.7 * cosine_score + 0.3 * fuzzy_score
                            if combined_score > best_combined:
                                best_combined = combined_score
                                best_detail = f"(Cr.P.C): {detail} (Combined Score: {combined_score:.2f})"
            if not candidate_found:
                # Instead of marking document as fake for missing sections, we can simply skip.
                continue
            else:
                valid_candidates.append((sec, best_combined, best_detail))
                all_combined_scores.append(best_combined)

        if valid_candidates:
            print(f"\n✅ Processed sentence:\n\"{sentence.strip()}\"")
            for candidate in sorted(valid_candidates, key=lambda x: int(x[0])):
                print(f"Section {candidate[0]}: {candidate[2]}")
        elif extracted_sections:
            print(f"\n⚠️ In sentence:\n\"{sentence.strip()}\"")
            print("No sections could be processed.")

    # ---------- Final Document Classification ----------
    print("\n========== Overall Document Citation Analysis ==========")
    if not all_combined_scores:
        print("No valid citation scores computed. (No sections found; assuming document is authentic.)")
        overall_avg = None
        document_fake = False
    else:
        overall_avg = sum(all_combined_scores) / len(all_combined_scores)
        print(f"Average Combined Similarity Score: {overall_avg:.2f}")
        if overall_avg < OVERALL_THRESHOLD:
            document_fake = True

    if document_fake:
        print("\n🚨 Document is FAKE due to incorrect section numbers and low overall similarity scores!")
    else:
        print("\n✅ Document appears to have authentic citation usage.")
    return not document_fake

#############################################
# 1. Load Articles from Final_IC.csv
#############################################
def load_articles_from_csv(csv_path):
    """
    Expects a CSV with columns:
      - article_id  (e.g. "Article 1 of Indian Constitution")
      - article_desc (the text/definition)
    Returns a dictionary: { article_id: article_desc, ... }
    """
    df = pd.read_csv(csv_path)
    article_dict = {}
    for _, row in df.iterrows():
        key = str(row['article_id']).strip()
        val = str(row['article_desc']).strip()
        article_dict[key] = val
    return article_dict

#############################################
# 2. Text Extraction (PDF or DOCX)
#############################################
def extract_text_from_file(file_path):
    """
    Extracts text from PDF or DOCX.
    For PDF, uses PyPDF2 (reads all pages).
    For DOCX, uses python-docx.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    elif ext in [".docx", ".doc"]:
        doc = Document(file_path)
        return "\n".join(para.text for para in doc.paragraphs)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def preprocess_text(text):
    """
    Lowercase + condense whitespace.
    """
    text = text.lower()
    text = re.sub(r'\s+', ' ', text).strip()
    return text

#############################################
# 3. Similarity Functions
#############################################
def fuzzy_similarity(str1, str2):
    """Compute a ratio [0..1] of textual similarity using difflib."""
    return difflib.SequenceMatcher(None, str1, str2).ratio()

def check_article_context(sentence, article_key, article_dict, st_model):
    """
    For a given sentence and an article key (like "Article 1 of Indian Constitution"),
    retrieve the reference text from `article_dict` and compute:
      - Cosine similarity with SentenceTransformer
      - Fuzzy similarity with difflib
    Returns (cosine_score, fuzzy_score).
    """
    reference_text = article_dict.get(article_key, "")
    if not reference_text:
        return None, None
    sentence_emb = st_model.encode(sentence, convert_to_tensor=True)
    ref_emb = st_model.encode(reference_text, convert_to_tensor=True)
    cosine_score = util.pytorch_cos_sim(sentence_emb, ref_emb).item()
    fuzzy_score = fuzzy_similarity(sentence, reference_text)
    return cosine_score, fuzzy_score

#############################################
# 4. Main Script: Checking Articles in Document
#############################################
def articles_checker(file_path):
    """
    Processes the given PDF or DOCX file, scanning its text for article references.
    Returns True (document authentic) if no discrepancies are found.
    If no article references are found, the document is assumed authentic.
    """
    # 4a. Load articles from CSV
    csv_path = "Final_IC.csv"  # <-- update with your CSV file path
    article_dict = load_articles_from_csv(csv_path)

    # 4b. Load SentenceTransformer model
    try:
        st_model = SentenceTransformer('nlpaueb/legal-bert-base-uncased')
        #print("Loaded legal-specific model (legal-bert-base-uncased).")
    except Exception as e:
        #print("Falling back to a general-purpose model. Error loading legal-bert:", e)
        st_model = SentenceTransformer('all-MiniLM-L6-v2')

    SIMILARITY_THRESHOLD = 0.60  # Adjust as needed

    # 4c. Extract text from your PDF or DOCX
    uploaded_file = file_path  # or .docx
    text = extract_text_from_file(uploaded_file)
    text = preprocess_text(text)
    document_fake = False
    articles_found = False  # Flag to track if any article references are found

    # 4d. Split text into sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)

    # 4e. Regex pattern to find references like "Article 1", "Article 2A", etc.
    article_pattern = re.compile(r'\barticle\s+(\d+[a-z]?)\b', re.IGNORECASE)

    # 4f. Scan each sentence for article references
    for sentence in sentences:
        found_articles = article_pattern.findall(sentence)
        if not found_articles:
            continue

        articles_found = True
        for art_num in found_articles:
            article_key = f"Article {art_num} of Indian Constitution"
            if article_key not in article_dict:
                print(f"\nSentence:\n\"{sentence.strip()}\"")
                print(f"Article {art_num} not found in dictionary keys.")
                continue

            cos_score, fz_score = check_article_context(sentence, article_key, article_dict, st_model)
            if cos_score is not None and fz_score is not None:
                combined_score = 0.7 * cos_score + 0.3 * fz_score
                print(f"\nSentence:\n\"{sentence.strip()}\"")
                print(f"Found {article_key}, Combined Score: {combined_score:.2f} (cos={cos_score:.2f}, fuzzy={fz_score:.2f})")
                if combined_score < SIMILARITY_THRESHOLD:
                    print("There should be an article mismatch. Fake detected.")
                    document_fake = True
                else:
                    print("  -> Looks good.")

    if not articles_found:
        print("\nNo article references found; assuming document is authentic.")
        document_fake = False

    return not document_fake

def verify_document(file_path):
    original_file=extract_text_from_file("RTDoc-12.pdf")

    # Execute all verification steps
    a=verify_document_structure(file_path,"RTDoc-12.pdf")
    b=verify_case_information(file_path)
    c=verify_document_date(file_path)
    d=load_legal_resources()
    e=section_checker(file_path)
    f=articles_checker(file_path)
    is_fake=a and b and c and d and e and f
    # Final result.
    if not is_fake:
        print(f"\n🚨 {file_path} is FLAGGED as FAKE.")
    else:
        print(f"\n✅ {file_path} appears AUTHENTIC.")

    return not is_fake

import os
import pandas as pd

# Global storage to accumulate batch-wise results
batch_results = {
    "correct_predictions": 0,
    "total_documents": 0
}

def load_inventory(excel_file):
    df = pd.read_excel(excel_file)
    inventory = {row['Filename']: row['Is_Fake'] for _, row in df.iterrows()}
    return inventory

def chunkify(lst, n):
    """Split list into n (almost equal) chunks"""
    avg = len(lst) // n
    remainder = len(lst) % n
    chunks = []
    start = 0
    for i in range(n):
        end = start + avg + (1 if i < remainder else 0)
        chunks.append(lst[start:end])
        start = end
    return chunks

# Get inputs from user
folder_path = input("Enter folder path containing documents: ").strip()
inventory_file = input("Enter inventory Excel file path: ").strip()

# Load all file names from folder that end with .pdf or .docx
all_files = [f for f in os.listdir(folder_path) if f.lower().endswith(('.pdf', '.docx'))]
batches = chunkify(all_files, 6)  # Split into 6 batches

# Load inventory once
inventory = load_inventory(inventory_file)

def run_batch(batch_index):
    batch_files = batches[batch_index]
    print(f"\n🚀 Running batch {batch_index + 1} with {len(batch_files)} files...")

    correct = 0
    for filename in batch_files:
        file_path = os.path.join(folder_path, filename)
        print(f"\n📄 Verifying: {filename}")

        try:
            # Assuming verify_document is defined elsewhere
            is_fake_detected = verify_document(file_path)
            detected_label = 1 if is_fake_detected else 0
            base_filename = os.path.splitext(filename)[0]
            ground_truth = inventory.get(base_filename, -1)
            if ground_truth == -1:
                print(f"⚠️ No ground truth for {filename}, skipping.")
                continue

            if detected_label == ground_truth:
                correct += 1
                print("✅ Correct classification!")
            else:
                print("❌ Incorrect classification!")
        except Exception as e:
            print(f"❌ Error verifying {filename}: {e}")

    batch_results["correct_predictions"] += correct
    batch_results["total_documents"] += len(batch_files)
    accuracy = (correct / len(batch_files)) * 100 if batch_files else 0
    print(f"\n📊 Batch {batch_index + 1} Accuracy: {accuracy:.2f}%")

# Run all batches using a for loop
for i in range(len(batches)):
    run_batch(i)

# Calculate and print the final overall accuracy
correct = batch_results["correct_predictions"]
total = batch_results["total_documents"]

if total > 0:
    final_accuracy = (correct / total) * 100
    print(f"\nFinal Accuracy across all batches: {final_accuracy:.2f}%")
    print(f"✅ Total Documents Processed: {total}")
    print(f"✅ Correctly Identified Documents: {correct}")
else:
    print("⚠️ No documents processed.")

"""# BERT"""

!pip install PyPDF2
import PyPDF2
import re
from collections import Counter

def generate_summary(text, num_sentences=3):
    sentences = text.split('. ')
    words = re.findall(r'\w+', text)
    word_frequencies = Counter(words)
    max_freq = max(word_frequencies.values(), default=1)

    sentence_scores = {}
    for sentence in sentences:
        for word in re.findall(r'\w+', sentence):
            if word in word_frequencies:
                sentence_scores[sentence] = sentence_scores.get(sentence, 0) + (word_frequencies[word] / max_freq)

    summary_sentences = sorted(sentence_scores, key=sentence_scores.get, reverse=True)[:num_sentences]
    return '. '.join(summary_sentences)


def preprocess_text(text):
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)  # Remove extra spaces
    return text.strip()

import collections
import unicodedata
import PyPDF2
import numpy as np
import re
from heapq import nlargest

# Custom Tokenization functions
def whitespace_tokenize(text):
    text = text.strip()
    if not text:
        return []
    return text.split()

def _is_whitespace(char):
    return char in [" ", "\t", "\n", "\r"] or unicodedata.category(char) == "Zs"

def _is_control(char):
    return unicodedata.category(char) in ("Cc", "Cf") and char not in ["\t", "\n", "\r"]

def _is_punctuation(char):
    cp = ord(char)
    return (33 <= cp <= 47) or (58 <= cp <= 64) or (91 <= cp <= 96) or (123 <= cp <= 126) or unicodedata.category(char).startswith("P")

class CustomTokenizer:
    def __init__(self, do_lower_case=True):
        self.do_lower_case = do_lower_case

    def tokenize(self, text):
        text = self._clean_text(text)
        tokens = whitespace_tokenize(text)
        split_tokens = []
        for token in tokens:
            if self.do_lower_case:
                token = token.lower()
            split_tokens.extend(self._split_on_punctuation(token))
        return whitespace_tokenize(" ".join(split_tokens))

    def _split_on_punctuation(self, text):
        chars = list(text)
        output = []
        current_token = []
        for char in chars:
            if _is_punctuation(char):
                if current_token:
                    output.append("".join(current_token))
                    current_token = []
                output.append(char)
            else:
                current_token.append(char)
        if current_token:
            output.append("".join(current_token))
        return output

    def _clean_text(self, text):
        return "".join(char if not _is_control(char) else " " for char in text)

# Extract text from PDF
def extract_text_from_pdf(pdf_file):
    with open(pdf_file, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        text = " ".join(page.extract_text() or "" for page in reader.pages)
    return text

# Generate summary without prebuilt models
def generate_summary(text, num_sentences=3):
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', text)
    word_frequencies = {}
    tokenizer = CustomTokenizer(do_lower_case=True)
    words = tokenizer.tokenize(text)

    for word in words:
        word_frequencies[word] = word_frequencies.get(word, 0) + 1

    max_freq = max(word_frequencies.values(), default=1)
    for word in word_frequencies:
        word_frequencies[word] /= max_freq

    sentence_scores = {}
    for sentence in sentences:
        for word in tokenizer.tokenize(sentence):
            if word in word_frequencies:
                sentence_scores[sentence] = sentence_scores.get(sentence, 0) + word_frequencies[word]

    summary_sentences = nlargest(num_sentences, sentence_scores, key=sentence_scores.get)
    return " ".join(summary_sentences)

# Main function
def main(input_pdf, output_summary_file):
    text = extract_text_from_pdf(input_pdf)
    tokenizer = CustomTokenizer(do_lower_case=True)
    tokens = tokenizer.tokenize(text)
    vocab = set(tokens)

    vocab_file = "vocab.txt"
    with open(vocab_file, 'w') as f:
        for token in sorted(vocab):
            f.write(token + "\n")

    summary = generate_summary(text)

    # Write summary to a file
    with open(output_summary_file, 'w') as f:
        f.write(summary)

    print(f"Generated Summary saved to {output_summary_file}")
    print(f"Vocabulary saved to {vocab_file}")
    print("Summary:\n", summary)

# Run the process
input_pdf = "Document (44).pdf"
output_summary_file = "summary-1.txt"
main(input_pdf, output_summary_file)

import nltk
import re
! pip install rouge_score bert_score sumy
from rouge_score import rouge_scorer
from bert_score import score
from sklearn.metrics import f1_score, recall_score
from sumy.parsers.plaintext import PlaintextParser
from sumy.summarizers.text_rank import TextRankSummarizer
from sumy.nlp.tokenizers import Tokenizer

# Ensure necessary nltk resources are available
nltk.download('punkt_tab')
nltk.download('punkt')
nltk.download('stopwords')
from nltk.corpus import stopwords

STOPWORDS = set(stopwords.words("english"))

# Function to clean text (lowercase, remove stopwords, punctuation)
def preprocess_text(text):
    text = text.lower().strip()
    text = re.sub(r'\s+', ' ', text)  # Normalize spaces
    text = re.sub(r'[^\w\s]', '', text)  # Remove punctuation
    text = " ".join([word for word in text.split() if word not in STOPWORDS])  # Remove stopwords
    return text

# Function to generate reference summary using TextRank
def create_reference_baseline(text):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = TextRankSummarizer()

    num_sentences = max(5, int(len(text.split(".")) * 0.3))  # Extract ~30% of sentences
    summary_sentences = summarizer(parser.document, num_sentences)

    reference_summary = " ".join([str(sentence) for sentence in summary_sentences])

    with open("reference.txt", "w") as ref_file:
        ref_file.write(reference_summary)

    return "reference.txt"

# Function to calculate ROUGE scores
def calculate_rouge(reference, summary):
    scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)
    scores = scorer.score(preprocess_text(reference), preprocess_text(summary))

    return {
        "ROUGE-1": {
            "Precision": round(scores['rouge1'].precision, 4),
            "Recall": round(scores['rouge1'].recall, 4),
            "F1": round(scores['rouge1'].fmeasure, 4)
        },
        "ROUGE-2": {
            "Precision": round(scores['rouge2'].precision, 4),
            "Recall": round(scores['rouge2'].recall, 4),
            "F1": round(scores['rouge2'].fmeasure, 4)
        },
        "ROUGE-L": {
            "Precision": round(scores['rougeL'].precision, 4),
            "Recall": round(scores['rougeL'].recall, 4),
            "F1": round(scores['rougeL'].fmeasure, 4)
        }
    }

# Function to calculate BERTScore
def calculate_bertscore(reference, summary):
    _, _, F1 = score([summary], [reference], lang="en", verbose=False)
    return round(F1.mean().item(), 4)  # Return only F1 score

# Function to calculate F1 Score and Recall manually
def calculate_f1_recall(reference, summary):
    ref_tokens = nltk.word_tokenize(preprocess_text(reference))
    sum_tokens = nltk.word_tokenize(preprocess_text(summary))

    ref_binary = [1 if token in sum_tokens else 0 for token in ref_tokens]
    sum_binary = [1 if token in ref_tokens else 0 for token in sum_tokens]

    max_len = max(len(ref_binary), len(sum_binary))
    ref_binary.extend([0] * (max_len - len(ref_binary)))
    sum_binary.extend([0] * (max_len - len(sum_binary)))

    return {
        "F1-Score": round(f1_score(ref_binary, sum_binary), 4),
        "Recall": round(recall_score(ref_binary, sum_binary), 4)
    }

def evaluate_summary(reference_file, summary_file):
    with open(reference_file, 'r') as ref_file:
        reference = ref_file.read().strip()
    with open(summary_file, 'r') as sum_file:
        summary = sum_file.read().strip()

    rouge_scores = calculate_rouge(reference, summary)
    bert_f1_score = calculate_bertscore(reference, summary)

    return rouge_scores, bert_f1_score


input_pdf = "Document (44).pdf"
text = extract_text_from_pdf(input_pdf)
reference_file = create_reference_baseline(text)
summary_file = "summary-1.txt"
rouge_results, bert_f1 = evaluate_summary(reference_file, summary_file)
print("\n===== ROUGE Evaluation =====\n")
for metric, scores in rouge_results.items():
    print(f"{metric}:")
    for sub_metric, score in scores.items():
        print(f"  {sub_metric}: {score}")
    print()

print("\n===== BERT F1 Score =====")
print(f"Overall BERT F1: {bert_f1}")

"""# BART"""

! pip install pdfminer.six
import os
import glob
import json
from collections import Counter
from pdfminer.high_level import extract_text

def extract_text_from_pdf(pdf_path):
    try:
        text = extract_text(pdf_path)
        return text
    except Exception as e:
        print(f"Error processing {pdf_path}: {e}")
        return ""

def build_vocab_from_pdfs(pdf_directory, vocab_size=10000, special_tokens=["<PAD>", "<SOS>", "<EOS>", "<MASK>"]):

    # Find all PDF files in the directory.
    pdf_files = glob.glob(os.path.join(pdf_directory, "*.pdf"))
    print(f"Found {len(pdf_files)} PDF files.")

    token_counter = Counter()

    # Process each PDF file.
    for pdf_file in pdf_files:
        print("Processing:", pdf_file)
        text = extract_text_from_pdf(pdf_file)
        tokens = text.lower().split()
        token_counter.update(tokens)

    vocab = {}
    for i, token in enumerate(special_tokens):
        vocab[token] = i

    num_special = len(special_tokens)
    most_common = token_counter.most_common(vocab_size - num_special)

    index = num_special
    for token, count in most_common:
        vocab[token] = index
        index += 1

    return vocab

def save_vocab(vocab, output_file):

    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(vocab, f, indent=2)
    print(f"Vocabulary saved to {output_file}")

if __name__ == "_main_":
    pdf_directory = "Documents"
    vocab_size = 10000
    output_vocab_file = "vocab.json"
    vocab = build_vocab_from_pdfs(pdf_directory, vocab_size=vocab_size)
    save_vocab(vocab, output_vocab_file)

import os
! pip install pdfplumber
import pdfplumber
import math
import copy
import torch
import torch.nn as nn
import torch.nn.functional as F
import torch.optim as optim
import random
import json

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

class PositionalEncoding(nn.Module):
    def __init__(self, d_model, dropout=0.1, max_len=5000):
        super().__init__()
        self.dropout = nn.Dropout(p=dropout)

        pe = torch.zeros(max_len, d_model)
        position = torch.arange(0, max_len, dtype=torch.float).unsqueeze(1)
        div_term = torch.exp(torch.arange(0, d_model, 2).float() * (-math.log(10000.0) / d_model))
        pe[:, 0::2] = torch.sin(position * div_term)
        pe[:, 1::2] = torch.cos(position * div_term)
        pe = pe.unsqueeze(0)
        self.register_buffer('pe', pe)

    def forward(self, x):
        if x.size(1) > self.pe.size(1):
            pe = torch.zeros(x.size(1), self.pe.size(2), device=x.device)
            position = torch.arange(0, x.size(1), dtype=torch.float, device=x.device).unsqueeze(1)
            div_term = torch.exp(torch.arange(0, self.pe.size(2), 2, device=x.device).float() * (-math.log(10000.0) / self.pe.size(2)))
            pe[:, 0::2] = torch.sin(position * div_term)
            pe[:, 1::2] = torch.cos(position * div_term)
            pe = pe.unsqueeze(0)
            x = x + pe[:, :x.size(1)]
        else:
            x = x + self.pe[:, :x.size(1)]
        return self.dropout(x)

class TransformerEncoderLayer(nn.Module):
    def __init__(self, d_model, nhead, dim_feedforward=2048, dropout=0.1):
        super().__init__()
        self.self_attn = nn.MultiheadAttention(d_model, nhead, dropout=dropout)
        self.linear1 = nn.Linear(d_model, dim_feedforward)
        self.dropout = nn.Dropout(dropout)
        self.linear2 = nn.Linear(dim_feedforward, d_model)
        self.norm1 = nn.LayerNorm(d_model)
        self.norm2 = nn.LayerNorm(d_model)
        self.dropout1 = nn.Dropout(dropout)
        self.dropout2 = nn.Dropout(dropout)

    def forward(self, src, src_mask=None, src_key_padding_mask=None):
        src2, _ = self.self_attn(src, src, src, attn_mask=src_mask, key_padding_mask=src_key_padding_mask)
        src = src + self.dropout1(src2)
        src = self.norm1(src)
        src2 = self.linear2(self.dropout(F.relu(self.linear1(src))))
        src = src + self.dropout2(src2)
        src = self.norm2(src)
        return src

class TransformerEncoder(nn.Module):
    def __init__(self, encoder_layer, num_layers):
        super().__init__()
        self.layers = nn.ModuleList([copy.deepcopy(encoder_layer) for _ in range(num_layers)])

    def forward(self, src, mask=None, src_key_padding_mask=None):
        output = src
        for mod in self.layers:
            output = mod(output, src_mask=mask, src_key_padding_mask=src_key_padding_mask)
        return output

class TransformerDecoderLayer(nn.Module):
    def __init__(self, d_model, nhead, dim_feedforward=2048, dropout=0.1):
        super().__init__()
        self.self_attn = nn.MultiheadAttention(d_model, nhead, dropout=dropout)
        self.multihead_attn = nn.MultiheadAttention(d_model, nhead, dropout=dropout)
        self.linear1 = nn.Linear(d_model, dim_feedforward)
        self.dropout = nn.Dropout(dropout)
        self.linear2 = nn.Linear(dim_feedforward, d_model)
        self.norm1 = nn.LayerNorm(d_model)
        self.norm2 = nn.LayerNorm(d_model)
        self.norm3 = nn.LayerNorm(d_model)
        self.dropout1 = nn.Dropout(dropout)
        self.dropout2 = nn.Dropout(dropout)
        self.dropout3 = nn.Dropout(dropout)

    def forward(self, tgt, memory, tgt_mask=None, memory_mask=None,
                tgt_key_padding_mask=None, memory_key_padding_mask=None):
        tgt2, _ = self.self_attn(tgt, tgt, tgt, attn_mask=tgt_mask, key_padding_mask=tgt_key_padding_mask)
        tgt = tgt + self.dropout1(tgt2)
        tgt = self.norm1(tgt)
        tgt2, _ = self.multihead_attn(tgt, memory, memory, attn_mask=memory_mask, key_padding_mask=memory_key_padding_mask)
        tgt = tgt + self.dropout2(tgt2)
        tgt = self.norm2(tgt)
        tgt2 = self.linear2(self.dropout(F.relu(self.linear1(tgt))))
        tgt = tgt + self.dropout3(tgt2)
        tgt = self.norm3(tgt)
        return tgt

class TransformerDecoder(nn.Module):
    def __init__(self, decoder_layer, num_layers):
        super().__init__()
        self.layers = nn.ModuleList([copy.deepcopy(decoder_layer) for _ in range(num_layers)])

    def forward(self, tgt, memory, tgt_mask=None, memory_mask=None,
                tgt_key_padding_mask=None, memory_key_padding_mask=None):
        output = tgt
        for mod in self.layers:
            output = mod(output, memory, tgt_mask=tgt_mask, memory_mask=memory_mask,
                         tgt_key_padding_mask=tgt_key_padding_mask,
                         memory_key_padding_mask=src_key_padding_mask if (src_key_padding_mask := None) else None)
        return output

class CustomBARTModel(nn.Module):
    def __init__(self, vocab_size, d_model=512, nhead=8,
                 num_encoder_layers=6, num_decoder_layers=6,
                 dim_feedforward=2048, dropout=0.1, max_seq_length=512):
        super().__init__()
        self.d_model = d_model
        self.embedding = nn.Embedding(vocab_size, d_model)
        self.pos_encoder = PositionalEncoding(d_model, dropout, max_seq_length)
        self.pos_decoder = PositionalEncoding(d_model, dropout, max_seq_length)
        encoder_layer = TransformerEncoderLayer(d_model, nhead, dim_feedforward, dropout)
        self.encoder = TransformerEncoder(encoder_layer, num_encoder_layers)
        decoder_layer = TransformerDecoderLayer(d_model, nhead, dim_feedforward, dropout)
        self.decoder = TransformerDecoder(decoder_layer, num_decoder_layers)
        self.output_projection = nn.Linear(d_model, vocab_size)

    def generate_square_subsequent_mask(self, sz):
        mask = (torch.triu(torch.ones(sz, sz)) == 1).transpose(0, 1)
        mask = mask.float().masked_fill(mask == 0, float('-inf')).masked_fill(mask == 1, float(0.0))
        return mask

    def forward(self, src, tgt, src_mask=None, tgt_mask=None,
                src_key_padding_mask=None, tgt_key_padding_mask=None):
        src_emb = self.embedding(src) * math.sqrt(self.d_model)
        src_emb = self.pos_encoder(src_emb)
        src_emb = src_emb.transpose(0, 1)
        memory = self.encoder(src_emb, mask=src_mask, src_key_padding_mask=src_key_padding_mask)
        tgt_emb = self.embedding(tgt) * math.sqrt(self.d_model)
        tgt_emb = self.pos_decoder(tgt_emb)
        tgt_emb = tgt_emb.transpose(0, 1)
        if tgt_mask is None:
            tgt_seq_len = tgt_emb.size(0)
            tgt_mask = self.generate_square_subsequent_mask(tgt_seq_len).to(tgt_emb.device)
        output = self.decoder(tgt_emb, memory, tgt_mask=tgt_mask, memory_mask=src_mask,
                              tgt_key_padding_mask=tgt_key_padding_mask,
                              memory_key_padding_mask=src_key_padding_mask)
        output = output.transpose(0, 1)
        logits = self.output_projection(output)
        return logits

with open("vocab.json", "r", encoding="utf-8") as f:
    loaded_vocab = json.load(f)
vocab = loaded_vocab
rev_vocab = {int(index): token for token, index in vocab.items()}

def simple_tokenizer(text):
    tokens = text.lower().split()
    token_ids = [vocab.get(token, vocab.get("this", 0)) for token in tokens]
    return token_ids

def simple_detokenizer(token_ids):
    words = [rev_vocab.get(token, "<UNK>") for token in token_ids]
    return " ".join(words)

def add_noise_to_input(input_tensor, mask_prob=0.3):
    noisy = input_tensor.clone()
    special_tokens = {vocab["<SOS>"], vocab["<EOS>"], vocab["<PAD>"]}
    mask_token = vocab["<MASK>"]
    for i in range(noisy.shape[0]):
        for j in range(noisy.shape[1]):
            if noisy[i, j].item() not in special_tokens:
                if random.random() < mask_prob:
                    noisy[i, j] = mask_token
    return noisy

def extract_text_from_pdfs(directory):
    documents = []
    for filename in os.listdir(directory):
        if filename.endswith('.pdf'):
            file_path = os.path.join(directory, filename)
            with pdfplumber.open(file_path) as pdf:
                text = ''
                for page in pdf.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += extracted_text + ' '  # Extract text from each page
                documents.append(text)
    return documents

pdf_directory = 'Documents'  # CHANGE this to your folder path
documents = extract_text_from_pdfs(pdf_directory)

# Model parameters
vocab_size = max(vocab.values()) + 1
d_model = 512
nhead = 8
num_encoder_layers = 6
num_decoder_layers = 6
dim_feedforward = 2048
dropout = 0.1
max_seq_length = 128

model = CustomBARTModel(vocab_size, d_model, nhead, num_encoder_layers, num_decoder_layers,
                        dim_feedforward, dropout, max_seq_length).to(device)
optimizer = optim.Adam(model.parameters(), lr=0.0001)
criterion = nn.CrossEntropyLoss(ignore_index=vocab["<PAD>"])

num_epochs = 15
model.train()
for epoch in range(num_epochs):
    total_loss = 0.0
    for doc in documents:
        doc_token_ids = simple_tokenizer(doc)
        if len(doc_token_ids) == 0:
            continue  # Skip empty documents
        if len(doc_token_ids) > max_seq_length - 1:
            doc_token_ids = doc_token_ids[:max_seq_length - 1]
        if doc_token_ids[-1] != vocab["<EOS>"]:
            doc_token_ids.append(vocab["<EOS>"])
        input_tensor = torch.tensor(doc_token_ids).unsqueeze(0).to(device)
        noisy_input = add_noise_to_input(input_tensor, mask_prob=0.2)
        decoder_input = torch.cat([
            torch.tensor([[vocab["<SOS>"]]]).to(device),
            input_tensor[:, :-1]
        ], dim=1)
        optimizer.zero_grad()
        output_logits = model(noisy_input, decoder_input)
        loss = criterion(output_logits.view(-1, vocab_size), input_tensor.view(-1))
        loss.backward()
        torch.nn.utils.clip_grad_norm_(model.parameters(), max_norm=1.0)
        optimizer.step()
        total_loss += loss.item()
    avg_loss = total_loss / len(documents)
    print(f"Epoch {epoch+1}, Loss: {avg_loss:.4f}")

def summarize_text(text, model, max_len=150, repetition_penalty=2.0, top_k=30, temperature=0.8):
    """
    Generate a summary using top-k sampling with repetition penalty.
    """
    model.eval()
    token_ids = simple_tokenizer(text)
    if len(token_ids) == 0:
        return ""
    if len(token_ids) > max_seq_length - 1:
        token_ids = token_ids[:max_seq_length - 1]
    if token_ids[-1] != vocab["<EOS>"]:
        token_ids.append(vocab["<EOS>"])
    input_tensor = torch.tensor(token_ids).unsqueeze(0).to(device)

    with torch.no_grad():
        src_emb = model.embedding(input_tensor) * math.sqrt(model.d_model)
        src_emb = model.pos_encoder(src_emb)
        src_emb = src_emb.transpose(0, 1)
        memory = model.encoder(src_emb)

        generated = [vocab["<SOS>"]]
        for _ in range(max_len):
            current_input = torch.tensor(generated).unsqueeze(0).to(device)
            tgt_emb = model.embedding(current_input) * math.sqrt(model.d_model)
            tgt_emb = model.pos_decoder(tgt_emb)
            tgt_emb = tgt_emb.transpose(0, 1)
            tgt_mask = model.generate_square_subsequent_mask(tgt_emb.size(0)).to(device)
            out = model.decoder(tgt_emb, memory, tgt_mask=tgt_mask)
            out = out.transpose(0, 1)
            logits = model.output_projection(out)  # shape: (1, current_length, vocab_size)
            next_token_logits = logits[0, -1, :].clone()
            next_token_logits = next_token_logits / temperature

            for token in set(generated):
                if token in {vocab["<SOS>"], vocab["<EOS>"], vocab["<PAD>"]}:
                    continue
                if next_token_logits[token] > 0:
                    next_token_logits[token] /= repetition_penalty
                else:
                    next_token_logits[token] *= repetition_penalty

            values, indices = torch.topk(next_token_logits, top_k)
            probabilities = F.softmax(values, dim=-1)
            sampled_index = torch.multinomial(probabilities, 1).item()
            next_token = indices[sampled_index].item()
            generated.append(next_token)
            if next_token == vocab["<EOS>"]:
                break

    summary_tokens = [token for token in generated[1:] if token != vocab["<EOS>"]]
    summary = simple_detokenizer(summary_tokens)
    return summary

def beam_search_summarize_text(text, model, max_len=200, beam_size=7, repetition_penalty=2.5, temperature=0.7):
    model.eval()
    token_ids = simple_tokenizer(text)
    if len(token_ids) == 0:
        return ""
    if len(token_ids) > max_seq_length - 1:
        token_ids = token_ids[:max_seq_length - 1]
    if token_ids[-1] != vocab["<EOS>"]:
        token_ids.append(vocab["<EOS>"])
    input_tensor = torch.tensor(token_ids).unsqueeze(0).to(device)

    with torch.no_grad():
        src_emb = model.embedding(input_tensor) * math.sqrt(model.d_model)
        src_emb = model.pos_encoder(src_emb)
        src_emb = src_emb.transpose(0, 1)
        memory = model.encoder(src_emb)

        beams = [([vocab["<SOS>"]], 0.0)]
        for _ in range(max_len):
            new_beams = []
            for seq, score in beams:
                if seq[-1] == vocab["<EOS>"]:
                    new_beams.append((seq, score))
                    continue
                current_input = torch.tensor(seq).unsqueeze(0).to(device)
                tgt_emb = model.embedding(current_input) * math.sqrt(model.d_model)
                tgt_emb = model.pos_decoder(tgt_emb)
                tgt_emb = tgt_emb.transpose(0, 1)
                tgt_mask = model.generate_square_subsequent_mask(tgt_emb.size(0)).to(device)
                out = model.decoder(tgt_emb, memory, tgt_mask=tgt_mask)
                out = out.transpose(0, 1)
                logits = model.output_projection(out)  # (1, seq_len, vocab_size)
                next_token_logits = logits[0, -1, :].clone()
                next_token_logits = next_token_logits / temperature

                # Apply repetition penalty
                for token in set(seq):
                    if token in {vocab["<SOS>"], vocab["<EOS>"], vocab["<PAD>"]}:
                        continue
                    if next_token_logits[token] > 0:
                        next_token_logits[token] /= repetition_penalty
                    else:
                        next_token_logits[token] *= repetition_penalty

                log_probs = F.log_softmax(next_token_logits, dim=-1)
                top_k = beam_size * 2  # Expand candidates for diversity
                values, indices = torch.topk(log_probs, top_k)

                for j in range(top_k):
                    token = indices[j].item()
                    new_score = score + values[j].item()
                    new_seq = seq + [token]
                    new_beams.append((new_seq, new_score))
            new_beams = sorted(new_beams, key=lambda x: x[1], reverse=True)[:beam_size]
            beams = new_beams
            if all(seq[-1] == vocab["<EOS>"] for seq, _ in beams):
                break

        best_seq, best_score = beams[0]
        final_seq = []
        for token in best_seq[1:]:
            if token == vocab["<EOS>"]:
                break
            final_seq.append(token)
        summary = simple_detokenizer(final_seq)
        return summary


if not os.path.exists(summary_folder):
    os.makedirs(summary_folder)

use_beam_search = True  # Set to False to use top-k sampling

for i, doc in enumerate(documents):
    if use_beam_search:
        summary = beam_search_summarize_text(doc, model, max_len=200, beam_size=5,
                                             repetition_penalty=2.0, temperature=0.8)
    else:
        summary = summarize_text(doc, model, max_len=200, repetition_penalty=2.0,
                                 top_k=30, temperature=0.8)

default_entries = [
        [r'\bpetition\b', 'request'],
        [r'\border date\b', 'order dated'],
        [r'\bfile\b', 'filed'],
        [r'\breject\b', 'rejected'],
        [r'\bprefer\b', 'seeks'],
        [r'\bclaim\b', 'alleged'],
        [r'\brental board\b', 'Rent Tribunal'],
        [r'\bimpugn pass\b', 'challenged'],
        [r'\binfirmity\b', 'defect'],
        [r'\billlegality\b', 'illegality'],
        [r'\bjurisdictional error\b', 'error in authority'],
        [r'\bexercise of its supervisory jurisdiction\b', 'its power to supervise']
    ]
csv_filename="dict.csv"
try:
    with open(csv_filename, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        for row in default_entries:
            writer.writerow(row)
    print(f"Default legal jargon CSV created as '{csv_filename}'.")
except Exception as e:
    print(f"Error creating CSV file {csv_filename}: {e}")

import os
import re
import heapq
import csv
import nltk
import spacy
! pip install pdfminer.six
! pip install language_tool_python
import language_tool_python
from pdfminer.high_level import extract_text
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords, wordnet as wn
from collections import Counter
from sklearn.feature_extraction.text import TfidfVectorizer

# Download necessary NLTK data
nltk.download('punkt')
nltk.download('punkt_tab')
nltk.download('stopwords')
nltk.download('wordnet')

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

def load_legal_jargon_dict(csv_filename):
    jargon_dict = {}
    try:
        with open(csv_filename, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            # Optionally, skip header row if present:
            # next(reader, None)
            for row in reader:
                if len(row) >= 2:
                    pattern = row[0].strip()
                    replacement = row[1].strip()
                    jargon_dict[pattern] = replacement
    except Exception as e:
        print(f"Error reading the CSV file {csv_filename}: {e}")
    return jargon_dict

def remove_metadata(text):
    # Define patterns for metadata fields (case-insensitive)
    metadata_patterns = [
        r'Request made by\s*:',
        r'Request made on\s*:',
        r'Client ID\s*:',
        r'Title\s*:',
        r'Delivery selection\s*:',
        r'Number of documents delivered\s*:',
        r'Delivery Summary'
    ]
    for pattern in metadata_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    return text

def preprocess_text(text):
    text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces/newlines with one space
    text = re.sub(r'[^\w\s\.,:;\-]', '', text)  # Keep letters, numbers, and common punctuation
    return text.strip()

def simplify_sentences(sentences):
    simplified = []
    for sent in sentences:
        doc = nlp(sent)
        tokens = []
        for token in doc:
            if token.is_punct:
                tokens.append(token.text)
            elif token.is_stop:
                tokens.append(token.text.lower())
            else:
                tokens.append(token.lemma_)
        simple_sent = " ".join(tokens)
        simple_sent = re.sub(r'\s([.,:;])', r'\1', simple_sent)  # Clean spacing around punctuation
        simplified.append(simple_sent)
    return ". ".join(simplified)

def get_important_sentences(text, keywords, num_sentences=5):
    sentences = sent_tokenize(text)
    sentence_scores = {}
    for sent in sentences:
        # Score each sentence by counting how many of the keywords appear in it.
        score = sum(sent.lower().count(word) for word in keywords)
        if score > 0:
            sentence_scores[sent] = score
    top_sentences = heapq.nlargest(num_sentences, sentence_scores, key=sentence_scores.get)
    # Preserve the original sentence order
    top_sentences = sorted(top_sentences, key=lambda s: sentences.index(s))
    return top_sentences

def improve_grammar(text):
    tool = language_tool_python.LanguageTool('en-US')
    corrected_text = tool.correct(text)
    return corrected_text

def reduce_legal_jargon_with_dict(text, jargon_dict):
    simplified_text = text
    for pattern, replacement in jargon_dict.items():
        simplified_text = re.sub(pattern, replacement, simplified_text, flags=re.IGNORECASE)
    return simplified_text

def context_aware_replacement(text, threshold=7):
    doc = nlp(text)
    new_tokens = []
    for token in doc:
        if token.is_alpha and len(token.text) > threshold and token.pos_ in ["ADJ", "ADV"]:
            new_tokens.append(get_simple_synonym(token))
        else:
            new_tokens.append(token.text)
    return " ".join(new_tokens)

def get_simple_synonym(token):
    pos_map = {"ADJ": wn.ADJ, "ADV": wn.ADV}
    if token.pos_ in pos_map:
        synsets = wn.synsets(token.text, pos=pos_map[token.pos_])
        if synsets:
            for syn in synsets:
                for lemma in syn.lemmas():
                    candidate = lemma.name().replace("_", " ")
                    if len(candidate) < len(token.text):
                        return candidate
    return token.text

import base64
import json
import zlib

def custom_encoder(input_data):

    try:
        json_str = json.dumps(input_data)
        compressed_data = zlib.compress(json_str.encode('utf-8'))
        encoded_str = base64.b64encode(compressed_data).decode('utf-8')
        return encoded_str
    except Exception as e:
        print(f"Encoding error: {e}")
        return None

def custom_decoder(encoded_str):
    try:
        compressed_data = base64.b64decode(encoded_str)
        json_str = zlib.decompress(compressed_data).decode('utf-8')
        data = json.loads(json_str)
        return data
    except Exception as e:
        print(f"Decoding error: {e}")
        return None

def rephrase_in_simple_english(text):
    return context_aware_replacement(text)

def train_tfidf_model(documents):
    vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
    vectorizer.fit(documents)
    return vectorizer

def extract_keywords_tfidf(text, vectorizer, top_n=20):
    tfidf_scores = vectorizer.transform([text])
    feature_names = vectorizer.get_feature_names_out()
    sorted_indices = tfidf_scores.toarray()[0].argsort()[-top_n:]
    keywords = [feature_names[i] for i in sorted_indices]
    return keywords

def extract_text_from_pdfs(directory):
    documents = []
    for filename in os.listdir(directory):
        if filename.lower().endswith(".pdf"):
            file_path = os.path.join(directory, filename)
            try:
                text = extract_text(file_path)
                if text:
                    documents.append(text)
            except Exception as e:
                print(f"Error reading {filename}: {e}")
    return documents

if __name__ == "__main__":
    jargon_csv = "dict.csv"
    legal_jargon_dict = load_legal_jargon_dict(jargon_csv)

    pdf_directory = 'Documents'
    training_documents = extract_text_from_pdfs(pdf_directory)
    if not training_documents:
        print("No training documents found in the specified folder.")
        exit(1)
    tfidf_vectorizer = train_tfidf_model(training_documents)

!apt install openjdk-17-jdk -y
!export JAVA_HOME=/usr/lib/jvm/java-17-openjdk-amd64
!export PATH=$PATH:$JAVA_HOME/bin
test_filename = "Document (44).pdf"
try:
    test_text = extract_text(test_filename)
except Exception as e:
    print(f"Error reading test document: {e}")
    test_text = ""

if test_text:
    cleaned_text = remove_metadata(test_text)
    cleaned_text = preprocess_text(cleaned_text)
    keywords = extract_keywords_tfidf(cleaned_text, tfidf_vectorizer, top_n=20)
    important_sentences = get_important_sentences(cleaned_text, keywords, num_sentences=5)
    raw_summary = simplify_sentences(important_sentences)
    improved_summary = improve_grammar(raw_summary)
    plain_summary = reduce_legal_jargon_with_dict(improved_summary, legal_jargon_dict)
    final_summary = rephrase_in_simple_english(plain_summary)
    base_name = os.path.splitext(os.path.basename(test_filename))[0]
    output_filename = f"summary-{base_name}.txt"
    try:
        with open(output_filename, "w", encoding="utf-8") as out_file:
            out_file.write("Simplified Summary in Plain English:\n")
            out_file.write(final_summary)
        print(f"Simplified summary saved to {output_filename}")
    except Exception as e:
        print(f"Error writing summary to file: {e}")
else:
    print("No text found in the test document.")

# Install necessary libraries:
! pip install bert-score sentence-transformers rake-nltk nltk

from bert_score import score
from sentence_transformers import SentenceTransformer, util
from rake_nltk import Rake
import nltk
nltk.download('punkt')
document = test_text
summary =final_summary
### Metric 1: BERTScore
P, R, F1 = score([summary], [document], lang="en", verbose=True)
print("BERTScore F1:",F1[0].item())

!pip install rouge-score

from rouge_score import rouge_scorer
scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)
rouge_scores = scorer.score(document, summary)
print("ROUGE Scores:")
for metric, score in rouge_scores.items():
    print(f"{metric.upper()}:")
    print(f"  Precision: {score.precision:.2f}")
    print(f"  Recall:    {score.recall:.2f}")
    print(f"  F-measure: {score.fmeasure:.2f}")

"""# Translation"""

import pandas as pd
import torch
import time
from torch.utils.data import Dataset, DataLoader
!pip install transformers[torch] --upgrade
from torch.optim import AdamW                  # <-- change here
from transformers import (
    MBartForConditionalGeneration,
    MBart50TokenizerFast
)
from nltk.translate.bleu_score import sentence_bleu

#######################################
# Utility Functions
#######################################
def fix_mojibake(text):
    """
    Fix common mojibake issues in text.
    """
    if not isinstance(text, str):
        return text
    if "à¤" in text:
        try:
            return text.encode("latin-1").decode("utf-8")
        except (UnicodeEncodeError, UnicodeDecodeError):
            return text
    return text

def clean_text_columns(df, columns):
    """
    Apply mojibake fixes, lowercasing, and trimming for specified columns.
    """
    for col in columns:
        df[col] = df[col].apply(fix_mojibake).str.lower().str.strip()
    return df

#######################################
# Section 1: English–Hindi Fine-Tuning & Evaluation
#######################################
def fine_tune_en_hi():
    # Load cleaned English–Hindi dataset (columns: "English", "Hindi")
    csv_file_path = "Cleaned_Corpus.csv"
    try:
        df = pd.read_csv(csv_file_path, encoding="utf-8-sig")
    except Exception as e:
        print(f"Error reading {csv_file_path}: {e}")
        return

    # Clean text columns
    df = clean_text_columns(df, ["English", "Hindi"])
    df["ref_hindi"] = df["Hindi"]

    # Initialize mBART50 model and tokenizer for many-to-many translation
    model_name = "facebook/mbart-large-50-many-to-many-mmt"
    tokenizer = MBart50TokenizerFast.from_pretrained(model_name)
    model = MBartForConditionalGeneration.from_pretrained(model_name)
    model.config.use_cache = False  # Reduce memory usage
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.to(device)

    # Set language codes for English → Hindi
    src_lang = "en_XX"
    tgt_lang = "hi_IN"
    tokenizer.src_lang = src_lang

    # Create a custom Dataset for translation fine-tuning
    class TranslationDataset(Dataset):
        def __init__(self, dataframe, tokenizer, max_length=128):
            self.dataframe = dataframe
            self.tokenizer = tokenizer
            self.max_length = max_length

        def __len__(self):
            return len(self.dataframe)

        def __getitem__(self, idx):
            source_text = self.dataframe.iloc[idx]["English"]
            target_text = self.dataframe.iloc[idx]["Hindi"]
            source_enc = self.tokenizer(
                source_text,
                max_length=self.max_length,
                truncation=True,
                padding="max_length",
                return_tensors="pt"
            )
            target_enc = self.tokenizer(
                target_text,
                max_length=self.max_length,
                truncation=True,
                padding="max_length",
                return_tensors="pt"
            )
            source_enc = {k: v.squeeze(0) for k, v in source_enc.items()}
            target_ids = target_enc["input_ids"].squeeze(0)
            source_enc["labels"] = target_ids
            return source_enc

    # Optionally use a subset for debugging
    # df = df.sample(n=100, random_state=42)

    train_dataset = TranslationDataset(df, tokenizer, max_length=128)
    train_dataloader = DataLoader(train_dataset, batch_size=4, shuffle=True)

    optimizer = AdamW(model.parameters(), lr=5e-5)
    num_epochs = 3

    # Simplified training loop
    model.train()
    for epoch in range(num_epochs):
        epoch_loss = 0.0
        for batch in train_dataloader:
            batch = {k: v.to(device) for k, v in batch.items()}
            optimizer.zero_grad()
            outputs = model(**batch)
            loss = outputs.loss
            loss.backward()
            optimizer.step()
            epoch_loss += loss.item()
        avg_loss = epoch_loss / len(train_dataloader)
        print(f"Epoch {epoch+1} completed. Average Loss: {avg_loss:.4f}")

    # Save the fine-tuned model and tokenizer
    model_save_path = "fine_tuned_mbart_en_hi"
    model.save_pretrained(model_save_path)
    tokenizer.save_pretrained(model_save_path)
    print(f"Fine-tuned model saved to '{model_save_path}'.")

    # Define translation function using the fine-tuned model
    def translate_en_to_hi(text):
        model.eval()
        inputs = tokenizer(text, return_tensors="pt", padding=True, truncation=True).to(device)
        generated_tokens = model.generate(
            **inputs,
            forced_bos_token_id=tokenizer.lang_code_to_id[tgt_lang]
        )
        return tokenizer.decode(generated_tokens[0], skip_special_tokens=True)

    # Translate and compute BLEU scores
    df["translated_hindi"] = df["English"].apply(translate_en_to_hi)
    def compute_bleu(reference, hypothesis):
        return sentence_bleu([reference.split()], hypothesis.split())
    df["bleu_score"] = df.apply(lambda row: compute_bleu(row["ref_hindi"], row["translated_hindi"]), axis=1)

    print("\nEnglish–Hindi Translation Samples:")
    print(df[["English", "ref_hindi", "translated_hindi", "bleu_score"]].head())

    comparison_csv_path = "Translated_Comparison_FineTuned.csv"
    df.to_csv(comparison_csv_path, index=False, encoding="utf-8-sig")
    print(f"Results saved as '{comparison_csv_path}'.")

#######################################
# Section 2: English–Tamil Translation & Evaluation
#######################################
def translate_en_ta():
    # Load cleaned English–Tamil dataset (columns: "en", "ta")
    csv_file_path = "Cleaned_English_Tamil_Dataset.csv"
    try:
        df = pd.read_csv(csv_file_path, encoding="utf-8-sig")
    except Exception as e:
        print(f"Error reading {csv_file_path}: {e}")
        return

    # Clean text for both English and Tamil
    df = clean_text_columns(df, ["en", "ta"])
    df["ref_tamil"] = df["ta"]

    # Optionally, use a smaller subset for testing/training
    num_records_to_keep = 10  # Adjust as needed
    df = df.head(num_records_to_keep)

    # Initialize the mBART50 model and tokenizer for translation
    model_name = "facebook/mbart-large-50-many-to-many-mmt"
    tokenizer = MBart50TokenizerFast.from_pretrained(model_name)
    model = MBartForConditionalGeneration.from_pretrained(model_name)
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.to(device)

    # Set language codes: English ("en_XX") to Tamil ("ta_IN")
    src_lang = "en_XX"
    tgt_lang = "ta_IN"
    tokenizer.src_lang = src_lang

    english_sentences = df["en"].tolist()
    batch_size = 8
    translations = []
    total_sentences = len(english_sentences)
    processed_count = 0
    last_print_time = time.time()

    # Process sentences in batches
    for i in range(0, total_sentences, batch_size):
        batch = english_sentences[i : i + batch_size]
        encoded = tokenizer(batch, return_tensors="pt", padding=True, truncation=True)
        encoded = {k: v.to(device) for k, v in encoded.items()}

        with torch.no_grad():
            generated_tokens = model.generate(
                **encoded,
                forced_bos_token_id=tokenizer.lang_code_to_id[tgt_lang]
            )
        batch_translations = tokenizer.batch_decode(generated_tokens, skip_special_tokens=True)
        translations.extend(batch_translations)
        processed_count += len(batch)
        current_time = time.time()
        if current_time - last_print_time >= 60:
            print(f"Processed {processed_count}/{total_sentences} sentences so far.")
            last_print_time = current_time

    print(f"Processed {processed_count}/{total_sentences} sentences in total.")
    df["translated_ta"] = translations

    # Compute BLEU scores for Tamil translation
    def compute_bleu(reference, hypothesis):
        return sentence_bleu([reference.split()], hypothesis.split())
    df["bleu_score"] = df.apply(lambda row: compute_bleu(row["ref_tamil"], row["translated_ta"]), axis=1)

    print("\nEnglish–Tamil Translation Samples:")
    print(df[["en", "ref_tamil", "translated_ta", "bleu_score"]].head())

    comparison_csv_path = "Translated_Comparison_Tamil.csv"
    df.to_csv(comparison_csv_path, index=False, encoding="utf-8-sig")
    print(f"Results saved as '{comparison_csv_path}'.")

#######################################
# Section 3: Translate a Text File to Hindi and Tamil
#######################################
def translate_text_file(file_path):
    """
    This function reads a text file line by line, translates each non-empty line
    from English to both Hindi and Tamil, and then writes the translations to separate files.
    """
    try:
        with open(file_path, 'r', encoding='utf-8-sig') as f:
            lines = f.read().splitlines()
    except Exception as e:
        print(f"Error reading the file {file_path}: {e}")
        return

    # Initialize model and tokenizer for both language translations.
    # If you have a fine-tuned model for Hindi, replace the model and tokenizer for Hindi below.
    model_name = "facebook/mbart-large-50-many-to-many-mmt"
    tokenizer_hi = MBart50TokenizerFast.from_pretrained(model_name)
    tokenizer_ta = MBart50TokenizerFast.from_pretrained(model_name)
    model_hi = MBartForConditionalGeneration.from_pretrained(model_name)
    model_ta = MBartForConditionalGeneration.from_pretrained(model_name)

    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model_hi.to(device)
    model_ta.to(device)

    # Set language codes for translation: English source
    src_lang = "en_XX"
    tgt_lang_hi = "hi_IN"
    tgt_lang_ta = "ta_IN"
    tokenizer_hi.src_lang = src_lang
    tokenizer_ta.src_lang = src_lang

    translations_hi = []
    translations_ta = []

    # Translate each non-empty line
    for line in lines:
        if line.strip():  # Only process non-empty lines
            # Hindi translation
            inputs_hi = tokenizer_hi(
                line, return_tensors="pt", padding=True, truncation=True
            ).to(device)
            with torch.no_grad():
                generated_hi = model_hi.generate(
                    **inputs_hi,
                    forced_bos_token_id=tokenizer_hi.lang_code_to_id[tgt_lang_hi]
                )
            translation_hi = tokenizer_hi.decode(generated_hi[0], skip_special_tokens=True)
            translations_hi.append(translation_hi)

            # Tamil translation
            inputs_ta = tokenizer_ta(
                line, return_tensors="pt", padding=True, truncation=True
            ).to(device)
            with torch.no_grad():
                generated_ta = model_ta.generate(
                    **inputs_ta,
                    forced_bos_token_id=tokenizer_ta.lang_code_to_id[tgt_lang_ta]
                )
            translation_ta = tokenizer_ta.decode(generated_ta[0], skip_special_tokens=True)
            translations_ta.append(translation_ta)
        else:
            translations_hi.append("")
            translations_ta.append("")

    # Save the translations to separate text files
    hi_output_path = "translated_hi.txt"
    ta_output_path = "translated_ta.txt"
    try:
        with open(hi_output_path, 'w', encoding='utf-8') as f_hi:
            for trans in translations_hi:
                f_hi.write(trans + "\n")
        with open(ta_output_path, 'w', encoding='utf-8') as f_ta:
            for trans in translations_ta:
                f_ta.write(trans + "\n")
        print(f"Translations saved to '{hi_output_path}' (Hindi) and '{ta_output_path}' (Tamil).")
    except Exception as e:
        print(f"Error writing translation files: {e}")

#######################################
# Main Execution
#######################################
if __name__ == "__main__":
    # Uncomment the following lines to run fine-tuning and evaluation.
    # print("Starting English–Hindi fine-tuning and evaluation...")
    # fine_tune_en_hi()
    #
    # print("\nStarting English–Tamil translation and evaluation...")
    # translate_en_ta()

    # For translating a text file, provide the path to your file.
    input_text_file = "input_text.txt"  # Update with the path to your text file
    print("\nTranslating text file to Hindi and Tamil...")
    translate_text_file(input_text_file)

"""# TEXT TO SPEECH

"""

!sudo apt-get update
!sudo apt-get install espeak-ng

import os
import subprocess
import regex
import numpy as np
import soundfile as sf
import unicodedata
import time
import librosa
import librosa.effects
from IPython.display import Audio, display

def get_espeak_executable():
    for cmd in ["espeak-ng", "espeak"]:
        try:
            path = subprocess.check_output(["which", cmd]).decode().strip()
            if path:
                return path
        except subprocess.CalledProcessError:
            continue
    return None

espeak_cmd = get_espeak_executable()
if espeak_cmd is None:
    exit(1)

def generate_phrase_audio(phrase, output_filename, language="en", voice_variant="en+f3"):
    command = [espeak_cmd, '-v', voice_variant, '-w', output_filename, phrase]
    try:
        subprocess.run(command, capture_output=True, text=True, check=True)
    except subprocess.CalledProcessError as e:
        print(f"Error generating audio for phrase '{phrase}':", e.stderr)
    time.sleep(0.5)  # Allow time for file writing
    if os.path.exists(output_filename) and os.path.getsize(output_filename) > 100:
        try:
            sr, data = sf.read(output_filename)
            return sr, data
        except Exception as ex:
            print(f"Error reading file {output_filename}: {ex}")
    print(f"Failed to synthesize audio for phrase '{phrase}'. Returning silence.")
    sr = 44100
    silence = np.zeros(int(0.1 * sr))
    return sr, silence

# Synthesize a Sentence at Phrase-Level
def synthesize_sentence_phraselevel(sentence, language="en", sample_rate=22050, phrase_pause=0.15):
    """
    Normalize and split the input sentence into phrases based on punctuation.
    Each phrase is synthesized as a whole using espeak-ng. A short pause is
    inserted between phrases to maintain natural breaks.

    Returns the final synthesized audio as a numpy array.
    """
    # Normalize sentence.
    sentence = unicodedata.normalize('NFC', sentence.strip())

    # Split the sentence into phrases. This regex splits after punctuation
    # (commas, periods, semicolons, exclamation and question marks) followed by whitespace.
    phrases = regex.split(r'(?<=[,.;?!])\s+', sentence)

    phrase_audio_dir = "phrase_audio"
    os.makedirs(phrase_audio_dir, exist_ok=True)
    segments = []

    # Synthesize each phrase.
    for phrase in phrases:
        if phrase:
            # Create a safe filename. Use a hash to avoid overly long filenames.
            safe_phrase = str(abs(hash(phrase)))
            output_filename = os.path.join(phrase_audio_dir, f"{safe_phrase}.wav")
            if not os.path.exists(output_filename):
                sr, data = generate_phrase_audio(phrase, output_filename, language, voice_variant="en+f3")
            if os.path.exists(output_filename) and os.path.getsize(output_filename) > 100:
                try:
                    audio, sr = sf.read(output_filename)
                    segments.append(audio)
                except Exception as e:
                    print(f"Error loading audio for phrase '{phrase}': {e}")
                    # Append a short pause if synthesis fails.
                    segments.append(np.zeros(int(phrase_pause * sample_rate)))
            else:
                segments.append(np.zeros(int(phrase_pause * sample_rate)))
            # Append a pause after each phrase.
            segments.append(np.zeros(int(phrase_pause * sample_rate)))

    # Concatenate all segments.
    if segments:
        synthesized = np.concatenate(segments)
        return synthesized
    else:
        return None

# Optional Prosody Modification (Speed and Pitch)
def adjust_speed(audio, speed_factor):
    """Time-stretch the audio using librosa."""
    return librosa.effects.time_stretch(audio.astype(np.float32), rate=speed_factor)

def adjust_pitch(audio, sr, n_steps):
    """Shift the pitch of the audio using librosa."""
    return librosa.effects.pitch_shift(audio.astype(np.float32), sr, n_steps)

# Synthesis, Post-Processing, and Playback
if __name__ == "__main__":
    # Read input text from the text file "input.txt"
    input_file_path = "input.txt"
    if not os.path.exists(input_file_path):
        print(f"Input file '{input_file_path}' not found. Please create the file and add your text.")
        exit(1)
    with open(input_file_path, "r", encoding="utf-8") as file:
        input_sentence = file.read()

    output_wav = "english_custom_phrase_tts.wav"

    # Synthesize at phrase-level.
    synthesized_audio = synthesize_sentence_phraselevel(
        input_sentence,
        language="en",
        sample_rate=22050,
        phrase_pause=0.5  # Adjust pause between phrases if needed.
    )
    if synthesized_audio is None:
        print("Could not synthesize audio from the input sentence.")
        exit(1)

    speed_factor = 0.5
    pitch_shift_steps = 0
    try:
        processed_audio = adjust_speed(synthesized_audio, speed_factor)
        processed_audio = adjust_pitch(processed_audio, 22050, pitch_shift_steps)
    except Exception as e:
        print("Error during prosody modification:", e)
        processed_audio = synthesized_audio

    # Save the final audio.
    sf.write(output_wav, processed_audio, 22050)
    print(f"Synthesized audio saved to {output_wav}")

    # Playback using IPython's Audio widget.
    display(Audio(output_wav, autoplay=True))

!pip install pdf2image pytesseract scikit-image
!apt-get install -y poppler-utils tesseract-ocr

import os
import subprocess
import regex
import numpy as np
import soundfile as sf
import unicodedata
import time
import librosa
import librosa.effects
import scipy.signal
from IPython.display import Audio, display

def get_espeak_executable():
    """Return the path to espeak-ng (or espeak if not available)."""
    for cmd in ["espeak-ng", "espeak"]:
        try:
            path = subprocess.check_output(["which", cmd]).decode().strip()
            if path:
                return path
        except subprocess.CalledProcessError:
            continue
    return None

espeak_cmd = get_espeak_executable()
if espeak_cmd is None:
    exit(1)

def generate_word_audio(word, output_filename, language="ta", voice_variant="ta+f4"):
    command = [espeak_cmd, '-v', voice_variant, '-w', output_filename, word]
    try:
        subprocess.run(command, capture_output=True, text=True, check=True)
    except subprocess.CalledProcessError as e:
        print(f"Error generating audio for word '{word}':", e.stderr)
    time.sleep(0.5)  # Allow time for the file to be written
    if os.path.exists(output_filename) and os.path.getsize(output_filename) > 100:
        try:
            sr, data = sf.read(output_filename)
            return sr, data
        except Exception as ex:
            print(f"Error reading file {output_filename}: {ex}")
    print(f"Failed to synthesize audio for word '{word}'. Returning silence.")
    sr = 44100
    silence = np.zeros(int(0.1 * sr))
    return sr, silence

def crossfade(audio1, audio2, fade_duration, sr):
    fade_samples = int(fade_duration * sr)
    fade_samples = min(fade_samples, len(audio1), len(audio2))
    fade_in = np.linspace(0, 1, fade_samples)
    fade_out = np.linspace(1, 0, fade_samples)
    audio1_fade = audio1.copy()
    audio2_fade = audio2.copy()
    audio1_fade[-fade_samples:] = audio1[-fade_samples:] * fade_out
    audio2_fade[:fade_samples] = audio2[:fade_samples] * fade_in
    return np.concatenate([audio1[:-fade_samples],
                           audio1_fade[-fade_samples:] + audio2_fade[:fade_samples],
                           audio2[fade_samples:]])

def synthesize_sentence_wordlevel(sentence, language="ta", sample_rate=22050, fade_duration=0.05):
    # Normalize sentence using NFC normalization.
    sentence = unicodedata.normalize('NFC', sentence)
    tokens = regex.split(r'(\s+)', sentence)

    word_audio_dir = "word_audio"
    os.makedirs(word_audio_dir, exist_ok=True)
    segments = []

    pause_short = np.zeros(int(0.05 * sample_rate))
    pause_long  = np.zeros(int(0.3 * sample_rate))
    default_fade_fraction = fade_duration

    for token in tokens:
        if token.isspace():
            segments.append(pause_short)
            continue
        token_stripped = token.strip()
        if token_stripped == "":
            segments.append(pause_short)
            continue

        # Determine pause duration based on punctuation.
        if token_stripped[-1] in {'.', ',', '!', '?'}:
            current_pause = pause_long
        else:
            current_pause = pause_short

        # Remove non-alphanumeric characters for a safe filename.
        safe_token = "".join(c for c in token_stripped if c.isalnum())
        if not safe_token:
            segments.append(current_pause)
            continue

        output_filename = os.path.join(word_audio_dir, f"{safe_token}.wav")
        if not os.path.exists(output_filename):
            sr, data = generate_word_audio(token_stripped, output_filename, language, voice_variant="ta+f4")
        if os.path.exists(output_filename) and os.path.getsize(output_filename) > 100:
            try:
                audio, sr = sf.read(output_filename)
                segments.append(audio)
                segments.append(current_pause)
            except Exception as e:
                print(f"Error loading audio for token '{token_stripped}': {e}")
                segments.append(current_pause)
        else:
            segments.append(current_pause)

    if segments:
        final_audio = segments[0]
        for seg in segments[1:]:
            final_audio = crossfade(final_audio, seg, default_fade_fraction, sample_rate)
        return final_audio
    else:
        return None

def adjust_speed(audio, speed_factor):
    """Time-stretch the audio using librosa."""
    return librosa.effects.time_stretch(audio.astype(np.float32), rate=speed_factor)

def adjust_pitch(audio, sr, n_steps):
    """Shift the pitch of the audio using librosa.effects.pitch_shift with keyword arguments."""
    return librosa.effects.pitch_shift(y=audio.astype(np.float32), sr=sr, n_steps=n_steps)

def normalize_audio(audio):
    """Normalize the amplitude of the audio using librosa.util.normalize."""
    return librosa.util.normalize(audio)

def apply_compression(audio, threshold=0.3, ratio=4.0):
    """
    Applies a simple dynamic range compressor.
    If the absolute amplitude exceeds the threshold, it is compressed according to the ratio.
    """
    abs_audio = np.abs(audio)
    sign_audio = np.sign(audio)
    compressed = np.where(abs_audio > threshold,
                          sign_audio * (threshold + (abs_audio - threshold) / ratio),
                          audio)
    return compressed

if __name__ == "__main__":
    # Read input text from the file "translated_ta.txt"
    input_file_path = "translated_ta.txt"
    if not os.path.exists(input_file_path):
        print(f"Input file '{input_file_path}' not found. Please create the file and add your Tamil text.")
        exit(1)
    with open(input_file_path, "r", encoding="utf-8") as file:
        input_sentence = file.read()

    output_wav = "tamil_custom_tts.wav"

    # Synthesize the sentence at the word-level with dynamic prosody and crossfade.
    synthesized_audio = synthesize_sentence_wordlevel(input_sentence, language="ta", sample_rate=22050, fade_duration=0.05)
    if synthesized_audio is None:
        print("Could not synthesize audio from the input sentence.")
        exit(1)

    # Optional prosody modifications.
    speed_factor = 1.1
    pitch_shift_steps = -1
    try:
        processed_audio = adjust_speed(synthesized_audio, speed_factor)
        processed_audio = adjust_pitch(processed_audio, 22050, pitch_shift_steps)
    except Exception as e:
        print("Error during prosody modification:", e)
        processed_audio = synthesized_audio

    # Additional effects.
    processed_audio = normalize_audio(processed_audio)
    processed_audio = apply_compression(processed_audio, threshold=0.3, ratio=4.0)

    # Save the final audio and playback.
    sf.write(output_wav, processed_audio, 22050)
    print(f"Synthesized audio saved to {output_wav}")
    display(Audio(output_wav, autoplay=True))

!pip install pdf2image pytesseract scikit-image
!apt-get install -y poppler-utils tesseract-ocr

import os
import subprocess
import regex
import numpy as np
import soundfile as sf
import unicodedata
import time
import librosa
import librosa.effects
import scipy.signal
from IPython.display import Audio, display

def get_espeak_executable():
    """Return the path to espeak-ng (or espeak if not available)."""
    for cmd in ["espeak-ng", "espeak"]:
        try:
            path = subprocess.check_output(["which", cmd]).decode().strip()
            if path:
                return path
        except subprocess.CalledProcessError:
            continue
    return None

espeak_cmd = get_espeak_executable()
if espeak_cmd is None:
    exit(1)

def generate_word_audio(word, output_filename, language="hi", voice_variant="hi+f3"):
    command = [espeak_cmd, '-v', voice_variant, '-w', output_filename, word]
    try:
        subprocess.run(command, capture_output=True, text=True, errors='replace', check=True)

    except subprocess.CalledProcessError as e:
        print(f"Error generating audio for word '{word}':", e.stderr)
    time.sleep(0.5)  # Allow time for the file to be written
    if os.path.exists(output_filename) and os.path.getsize(output_filename) > 100:
        try:
            sr, data = sf.read(output_filename)
            return sr, data
        except Exception as ex:
            print(f"Error reading file {output_filename}: {ex}")
    print(f"Failed to synthesize audio for word '{word}'. Returning silence.")
    sr = 44100
    silence = np.zeros(int(0.1 * sr))
    return sr, silence

def crossfade(audio1, audio2, fade_duration, sr):
    """
    Crossfade two 1D numpy audio arrays using a fade_duration in seconds.
    Returns the concatenated audio with an overlapping fade.
    """
    fade_samples = int(fade_duration * sr)
    fade_samples = min(fade_samples, len(audio1), len(audio2))
    fade_in = np.linspace(0, 1, fade_samples)
    fade_out = np.linspace(1, 0, fade_samples)
    audio1_fade = audio1.copy()
    audio2_fade = audio2.copy()
    audio1_fade[-fade_samples:] = audio1[-fade_samples:] * fade_out
    audio2_fade[:fade_samples] = audio2[:fade_samples] * fade_in
    return np.concatenate([audio1[:-fade_samples],
                           audio1_fade[-fade_samples:] + audio2_fade[:fade_samples],
                           audio2[fade_samples:]])

def synthesize_sentence_wordlevel(sentence, language="hi", sample_rate=22050, fade_duration=0.05):
    sentence = unicodedata.normalize('NFC', sentence)
    tokens = regex.split(r'(\s+)', sentence)

    word_audio_dir = "word_audio"
    os.makedirs(word_audio_dir, exist_ok=True)
    segments = []

    pause_short = np.zeros(int(0.05 * sample_rate))
    pause_long  = np.zeros(int(0.3 * sample_rate))
    default_fade_fraction = fade_duration

    for token in tokens:
        if token.isspace():
            segments.append(pause_short)
            continue
        token_stripped = token.strip()
        if token_stripped == "":
            segments.append(pause_short)
            continue

        # Determine pause duration based on punctuation.
        if token_stripped[-1] in {'.', ',', '!', '?'}:
            current_pause = pause_long
        else:
            current_pause = pause_short

        safe_token = "".join(c for c in token_stripped if c.isalnum())
        if not safe_token:
            segments.append(current_pause)
            continue

        output_filename = os.path.join(word_audio_dir, f"{safe_token}.wav")
        if not os.path.exists(output_filename):
            sr, data = generate_word_audio(token_stripped, output_filename, language, voice_variant="hi+f3")
        if os.path.exists(output_filename) and os.path.getsize(output_filename) > 100:
            try:
                audio, sr = sf.read(output_filename)
                segments.append(audio)
                segments.append(current_pause)
            except Exception as e:
                print(f"Error loading audio for token '{token_stripped}': {e}")
                segments.append(current_pause)
        else:
            segments.append(current_pause)

    if segments:
        final_audio = segments[0]
        for seg in segments[1:]:
            final_audio = crossfade(final_audio, seg, default_fade_fraction, sample_rate)
        return final_audio
    else:
        return None

def adjust_speed(audio, speed_factor):
    """Time-stretch the audio using librosa."""
    return librosa.effects.time_stretch(audio.astype(np.float32), rate=speed_factor)

def adjust_pitch(audio, sr, n_steps):
    """Shift the pitch of the audio using librosa.effects.pitch_shift."""
    return librosa.effects.pitch_shift(y=audio.astype(np.float32), sr=sr, n_steps=n_steps)

def normalize_audio(audio):
    """Normalize the amplitude of the audio using librosa.util.normalize."""
    return librosa.util.normalize(audio)

def apply_compression(audio, threshold=0.3, ratio=4.0):
    """
    Applies a simple dynamic range compressor.
    If the absolute amplitude exceeds the threshold, it is compressed according to the ratio.
    """
    abs_audio = np.abs(audio)
    sign_audio = np.sign(audio)
    compressed = np.where(abs_audio > threshold,
                          sign_audio * (threshold + (abs_audio - threshold) / ratio),
                          audio)
    return compressed

if __name__ == "__main__":
    # Read input text from the file "translated_hi.txt"
    input_file_path = "translated_hi.txt"
    if not os.path.exists(input_file_path):
        print(f"Input file '{input_file_path}' not found. Please create the file and add your Hindi text.")
        exit(1)
    with open(input_file_path, "r", encoding="utf-8") as file:
        input_sentence = file.read()

    output_wav = "hindi_custom_tts.wav"

    # Synthesize at word-level with dynamic prosody and crossfade.
    synthesized_audio = synthesize_sentence_wordlevel(input_sentence, language="hi", sample_rate=22050, fade_duration=0.05)
    if synthesized_audio is None:
        print("Could not synthesize audio from the input sentence.")
        exit(1)

    # Optional prosody modifications.
    speed_factor = 1.1
    pitch_shift_steps = -3
    try:
        processed_audio = adjust_speed(synthesized_audio, speed_factor)
        processed_audio = adjust_pitch(processed_audio, 22050, pitch_shift_steps)
    except Exception as e:
        print("Error during prosody modification:", e)
        processed_audio = synthesized_audio

    # Additional effects.
    processed_audio = normalize_audio(processed_audio)
    processed_audio = apply_compression(processed_audio, threshold=0.3, ratio=4.0)

    # Save and playback.
    sf.write(output_wav, processed_audio, 22050)
    print(f"Synthesized audio saved to {output_wav}")
    display(Audio(output_wav, autoplay=True))
